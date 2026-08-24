#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROMAN = {
    "I": 1,
    "II": 2,
    "III": 3,
    "IV": 4,
    "V": 5,
    "VI": 6,
    "VII": 7,
    "VIII": 8,
    "IX": 9,
    "X": 10,
}


def set_run_font(run, size: Pt = Pt(12)):
    run.font.name = "Times New Roman"
    run.font.size = size
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Times New Roman")
    rpr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rpr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def set_cell_width(cell, width_cm: float):
    width = Cm(width_cm)
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.twips)))
    tc_w.set(qn("w:type"), "dxa")


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "nil")


def chapter_number_from_heading(text: str) -> int | None:
    m = re.search(r"\bBAB\s+([IVXLCDM]+)\b", text.upper())
    if not m:
        return None
    return ROMAN.get(m.group(1))


def is_display_math(paragraph) -> bool:
    # Pandoc emits block TeX mathematics as an Office Math paragraph (m:oMathPara).
    # Inline math appears as m:oMath without m:oMathPara and is intentionally ignored.
    return bool(paragraph._p.xpath(".//m:oMathPara"))


def equation_table_for(paragraph, number: str, doc: Document):
    math_nodes = paragraph._p.xpath(".//m:oMathPara")
    if not math_nodes:
        raise RuntimeError("display-math paragraph lost its m:oMathPara node")

    table = doc.add_table(rows=1, cols=3)
    paragraph._p.addprevious(table._tbl)
    table.autofit = False
    remove_table_borders(table)

    # Portrait text area under SPs margins: 21 - 4 - 3 = 14 cm.
    # Keep a small balancing left gutter and a dedicated number column so the
    # equation remains visually centered while the number sits near the right margin.
    widths = (1.5, 10.5, 2.0)
    for cell, width in zip(table.rows[0].cells, widths):
        set_cell_width(cell, width)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    left_p = table.cell(0, 0).paragraphs[0]
    left_p.paragraph_format.first_line_indent = Cm(0)
    left_p.paragraph_format.space_before = Pt(0)
    left_p.paragraph_format.space_after = Pt(0)

    middle_p = table.cell(0, 1).paragraphs[0]
    middle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    middle_p.paragraph_format.first_line_indent = Cm(0)
    middle_p.paragraph_format.space_before = Pt(0)
    middle_p.paragraph_format.space_after = Pt(0)
    middle_p.paragraph_format.keep_together = True

    math_node = math_nodes[0]
    math_node.getparent().remove(math_node)
    middle_p._p.append(math_node)

    right_p = table.cell(0, 2).paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_p.paragraph_format.first_line_indent = Cm(0)
    right_p.paragraph_format.space_before = Pt(0)
    right_p.paragraph_format.space_after = Pt(0)
    run = right_p.add_run(number)
    set_run_font(run)

    paragraph._element.getparent().remove(paragraph._element)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--docx", type=Path, required=True)
    args = ap.parse_args()

    doc = Document(args.docx)
    chapter = None
    counters: dict[int, int] = {}
    numbered = 0

    # Snapshot paragraphs because display-math paragraphs are removed during conversion.
    for paragraph in list(doc.paragraphs):
        heading_chapter = chapter_number_from_heading(paragraph.text.strip())
        if paragraph.style.name == "Heading 1" and heading_chapter is not None:
            chapter = heading_chapter
            counters.setdefault(chapter, 0)
            continue

        if chapter is None or not is_display_math(paragraph):
            continue

        counters[chapter] = counters.get(chapter, 0) + 1
        label = f"({chapter}.{counters[chapter]})"
        equation_table_for(paragraph, label, doc)
        numbered += 1

    doc.save(args.docx)
    print(f"applied chapter-based equation numbering to {numbered} display equations")


if __name__ == "__main__":
    main()
