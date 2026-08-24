#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

TITLE = "ANALISIS DAN OPTIMASI PREPROCESSING CITRA BERBASIS FREKUENSI-ANGULAR PADA YOLO26 UNTUK DETEKSI FINE-GRAINED CACAT BIJI KOPI"

CITE = {
    "STD-01": "Badan Standardisasi Nasional, 2008",
    "REV-01": "Motta et al., 2024",
    "COF-01": "Hong et al., 2026",
    "COF-02": "Bahy & Rifai, 2026",
    "COF-03": "Samudra & Rachmawati, 2025",
    "COF-04": "Hebert & Alamsyah, 2026",
    "COF-05": "Jundullah et al., 2026",
    "COF-06": "Gope et al., 2024",
    "COF-07": "Kesiman et al., 2023",
    "COF-08": "Arwatchananukul et al., 2024",
    "COF-09": "Lei et al., 2025",
    "COF-10": "de Oliveira et al., 2016",
    "COF-11": "Chang & Huang, 2021",
    "COF-12": "Jiao et al., 2025",
    "COF-13": "Hu et al., 2025",
    "COF-14": "Muchtar et al., 2025",
    "COF-15": "Hsia et al., 2022",
    "COF-16": "Gope et al., 2025",
    "COF-17": "García et al., 2019",
    "DET-01": "Jocher et al., 2026",
    "DET-02": "Ren et al., 2015",
    "DET-03": "Redmon et al., 2016",
    "EVAL-01": "Lin et al., 2014",
    "EVAL-02": "COCOeval",
    "DIAG-01": "Feng et al., 2021",
    "DIAG-02": "Wu et al., 2020",
    "DIAG-03": "Jiang et al., 2018",
    "FG-01": "Xu et al., 2025",
    "FG-02": "Xie et al., 2025",
    "FG-03": "Wang et al., 2020",
    "PRE-01": "Liu et al., 2022",
    "PRE-02": "Qin et al., 2022",
    "PRE-03": "Li et al., 2025",
    "PRE-04": "Syauqi et al., 2025",
    "PRE-05": "Chen et al., 2024",
    "PRE-06": "Tu et al., 2026",
    "PRE-07": "Cai et al., 2023",
    "PRE-08": "Yang & Soatto, 2020",
    "SPEC-01": "Cao et al., 2019",
    "SPEC-02": "Zhang & Tan, 2003",
    "FREQ-01": "Chi et al., 2020",
    "FREQ-02": "Li et al., 2024",
    "FREQ-03": "Chen et al., 2025",
    "WAVE-01": "Finder et al., 2024",
    "THEORY-01": "Gonzalez & Woods",
    "THEORY-02": "Bracewell",
}

BIB = {
    "STD-01": "Badan Standardisasi Nasional. (2008). SNI 01-2907-2008: Biji kopi.",
    "COF-01": "Hong et al. (2026). Automated detection of defective coffee beans based on improved YOLOv10 framework. Current Research in Food Science. https://doi.org/10.1016/j.crfs.2026.101461",
    "COF-02": "Bahy, & Rifai. (2026). Real-time coffee bean defect detection based on SNI 01-2907-2008 standards using lightweight YOLOv5s architecture. International Journal on ICT. https://doi.org/10.21108/ijoict.v12i1.10584",
    "COF-03": "Samudra, & Rachmawati. (2025). Deep learning-based defect detection in Arabica green coffee beans using LSKNet. ICoDSA 2025.",
    "COF-04": "Hebert, & Alamsyah. (2026). Detection of coffee bean defects in Speciality Coffee Association standards using YOLOv12. INOVTEK Polbeng - Seri Informatika.",
    "COF-05": "Jundullah et al. (2026). YOLOv8-based multi-class detection of coffee bean defects and contaminants for automated quality grading. Brilliance: Research of Artificial Intelligence. https://doi.org/10.47709/brilliance.v6i2.8612",
    "COF-06": "Gope et al. (2024). Comparative analysis of YOLO models for green coffee bean detection and defect classification. Scientific Reports. https://doi.org/10.1038/s41598-024-78598-7",
    "COF-07": "Kesiman et al. (2023). Benchmarking a new dataset for coffee bean defects classification based on SNI 01-2907-2008. ICITRI 2023. https://doi.org/10.1109/ICITRI59340.2023.10249345",
    "COF-08": "Arwatchananukul et al. (2024). Implementing a deep learning model for defect classification in Thai Arabica green coffee beans. Smart Agricultural Technology. https://doi.org/10.1016/j.atech.2024.100680",
    "COF-09": "Lei et al. (2025). A coffee bean defect detection algorithm with decoupled classification and localization. ICIC 2025 / CCIS 2567. https://doi.org/10.1007/978-981-96-9952-0_26",
    "COF-10": "de Oliveira et al. (2016). A computer vision system for coffee beans classification based on computational intelligence techniques. Journal of Food Engineering. https://doi.org/10.1016/j.jfoodeng.2015.10.030",
    "COF-12": "Jiao et al. (2025). Swin-HSSAM: A green coffee bean grading method by Swin Transformer. PLOS ONE. https://doi.org/10.1371/journal.pone.0322198",
    "COF-13": "Hu et al. (2025). Siamese networks for few-shot coffee bean defect detection. LWT. https://doi.org/10.1016/j.lwt.2025.118631",
    "COF-14": "Muchtar et al. (2025). Edge AI-based detection for defective coffee beans using deep learning and Streamlit framework. IEEE Access. https://doi.org/10.1109/ACCESS.2025.3561189",
    "REV-01": "Motta et al. (2024). Machine learning techniques for coffee classification: A comprehensive review of scientific research. Artificial Intelligence Review.",
    "DET-01": "Jocher et al. (2026). Ultralytics YOLO26: Unified real-time end-to-end vision models. arXiv:2606.03748.",
    "DET-02": "Ren et al. (2015). Faster R-CNN: Towards real-time object detection with region proposal networks. NeurIPS 2015.",
    "DET-03": "Redmon et al. (2016). You only look once: Unified, real-time object detection. CVPR 2016.",
    "DIAG-01": "Feng et al. (2021). TOOD: Task-aligned one-stage object detection. ICCV 2021.",
    "DIAG-02": "Wu et al. (2020). Rethinking classification and localization for object detection. CVPR 2020.",
    "DIAG-03": "Jiang et al. (2018). Acquisition of localization confidence for accurate object detection. ECCV 2018.",
    "FG-01": "Xu et al. (2025). More signals matter to detection: Integrating language knowledge and frequency representations for boosting fine-grained aircraft recognition. Neural Networks, 187, 107402. https://doi.org/10.1016/j.neunet.2025.107402",
    "FG-02": "Xie et al. (2025). Learning discriminative representation for fine-grained object detection in remote sensing images. IEEE Transactions on Circuits and Systems for Video Technology.",
    "FG-03": "Wang et al. (2020). An adversarial domain adaptation network for cross-domain fine-grained recognition. WACV 2020.",
    "PRE-01": "Liu et al. (2022). Image-Adaptive YOLO for object detection in adverse weather conditions. AAAI 2022.",
    "PRE-02": "Qin et al. (2022). DENet: Detection-driven enhancement network for object detection under adverse weather conditions. ACCV 2022.",
    "PRE-03": "Li et al. (2025). FE-YOLO: Fourier enhancement YOLO for end-to-end object detection in low-light conditions. Digital Signal Processing.",
    "PRE-04": "Syauqi et al. (2025). Edge AI-based defect detection in white pepper using CLAHE-based preprocessing and YOLOv8. IEEE ICONS-IoT 2025.",
    "PRE-05": "Chen et al. (2024). Soft X-ray image recognition and classification of maize seed cracks based on image enhancement and YOLOv8. Computers and Electronics in Agriculture.",
    "PRE-06": "Tu et al. (2026). Wavelet quadtree contrast-limited adaptive histogram equalisation tablet enhancement method for defect detection. 2026.",
    "PRE-07": "Cai et al. (2023). Retinexformer: One-stage Retinex-based Transformer for low-light image enhancement. ICCV 2023.",
    "PRE-08": "Yang, & Soatto. (2020). FDA: Fourier domain adaptation for semantic segmentation. CVPR 2020.",
    "SPEC-01": "Cao et al. (2019). Frequency spectrum-based optimal texture window size selection for high spatial resolution remote sensing imagery. 2019.",
    "SPEC-02": "Zhang, & Tan. (2003). Affine invariant classification and retrieval of texture images. Pattern Recognition.",
    "FREQ-01": "Chi, Jiang, & Mu. (2020). Fast Fourier convolution. NeurIPS 2020.",
    "FREQ-02": "Li et al. (2024). FDADNet: Detection of surface defects in wood-based panels based on frequency domain transformation and adaptive dynamic downsampling. Processes. https://doi.org/10.3390/pr12102134",
    "FREQ-03": "Chen et al. (2025). Frequency dynamic convolution for dense image prediction. CVPR 2025.",
    "WAVE-01": "Finder et al. (2024). Wavelet convolutions for large receptive fields. ECCV 2024.",
    "EVAL-01": "Lin et al. (2014). Microsoft COCO: Common objects in context. ECCV 2014.",
    "EVAL-02": "COCO Consortium. COCOeval evaluation implementation/specification.",
}


def read(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def cut_between(text: str, start_heading: str, end_heading: str | None = None) -> str:
    start = text.find(start_heading)
    if start < 0:
        raise ValueError(f"heading not found: {start_heading}")
    out = text[start:]
    if end_heading:
        end = out.find(end_heading, len(start_heading))
        if end >= 0:
            out = out[:end]
    return out.strip()


def replace_section(base: str, start_heading: str, next_heading: str, replacement: str) -> str:
    a = base.find(start_heading)
    b = base.find(next_heading, a + len(start_heading))
    if a < 0 or b < 0:
        raise ValueError(f"cannot replace {start_heading}")
    return base[:a] + replacement.strip() + "\n\n---\n\n" + base[b:]


def strip_admin(text: str) -> str:
    lines = text.splitlines()
    out = []
    in_status_block = False
    for line in lines:
        if line.startswith("Status:") or line.startswith("> Status:"):
            continue
        if line.startswith("Purpose:") or line.startswith("Tujuan normalisasi:"):
            continue
        if line.startswith("All bracketed citation keys"):
            continue
        if line.startswith("This file") or line.startswith("Use this module"):
            continue
        out.append(line)
    return "\n".join(out).strip()


def replace_citations(text: str) -> tuple[str, set[str]]:
    used: set[str] = set()

    def repl(match: re.Match[str]) -> str:
        inside = match.group(1).strip()
        # leave markdown links and ordinary brackets alone
        keys = re.findall(r"(?:STD|REV|COF(?:-SUP)?|DET|EVAL|DIAG|FG|PRE|SPEC|FREQ|WAVE|THEORY)-\d+", inside)
        if not keys:
            return match.group(0)
        used.update(keys)
        # Locator form: [FG-01, p. 4, ...]
        if len(keys) == 1 and inside.startswith(keys[0]) and inside != keys[0]:
            tail = inside[len(keys[0]):].lstrip(" ,")
            return f"({CITE.get(keys[0], keys[0])}, {tail})" if tail else f"({CITE.get(keys[0], keys[0])})"
        authors = [CITE.get(k, k) for k in keys]
        return "(" + "; ".join(authors) + ")"

    return re.sub(r"\[([^\]\n]+)\]", repl, text), used


def demote_first_heading(text: str, old: str, new: str) -> str:
    return text.replace(old, new, 1)


def prepare_sources(repo: Path) -> tuple[str, set[str]]:
    proposal = repo / "docs/thesis/proposal"

    background = read(proposal / "02_BACKGROUND.md")
    background = cut_between(background, "# 1.1 Latar Belakang", "---\n\n## Citation-key map")
    background = demote_first_heading(background, "# 1.1", "## 1.1")

    problem = read(proposal / "03_PROBLEM_FORMULATION.md")
    problem = cut_between(problem, "## 1.2 Identifikasi Masalah", "---\n\n## Research logic")

    lit = read(proposal / "04_LITERATURE_REVIEW.md")
    lit = cut_between(lit, "## 2.1 Biji Kopi Hijau", None)
    normalized_22 = read(proposal / "04_02_INSPECTION_QUALITY_NORMALIZED.md")
    normalized_22 = cut_between(normalized_22, "# 2.2 Inspeksi", "## Source-role note")
    normalized_22 = demote_first_heading(normalized_22, "# 2.2", "## 2.2")
    related = read(proposal / "04_09_RELATED_WORK_TABLE.md")
    related = cut_between(related, "## Tabel 2.1", "## Anti-overclaim note")
    related = "## 2.9 Penelitian Terkait\n\n" + related
    lit = replace_section(lit, "## 2.2 Inspeksi", "## 2.3 Object Detection", normalized_22)
    lit = replace_section(lit, "## 2.9 Penelitian Terkait", "### Positioning summary", related)
    # Drop any legacy remainder after the normalized related-work section.
    pos = lit.find("### Positioning summary")
    if pos >= 0:
        lit = lit[:pos]

    method = read(proposal / "05_METHODOLOGY.md")
    method = cut_between(method, "## 3.1 Kerangka Penelitian", "---\n\n## Guardrails Bab III")
    hard35 = read(proposal / "05_05_AF2_PRIMARY_SOURCE_HARDENED.md")
    hard35 = cut_between(hard35, "## 3.5 Preprocessing Frekuensi-Angular AF2", None)
    # Remove obsolete open-locator paragraph from the hardened module.
    hard35 = re.sub(
        r"\*\*Catatan provenance:\*\* repository memberi anotasi eksplisit.*?gap definisi implementasi\.\n",
        "Xu et al. menjelaskan keluarga persamaan AFAB-2 pada pp. 5–6, §3.3.3, Eq. (9)–(13). Implementasi penelitian memetakan keluarga persamaan tersebut ke operator diskret yang dibekukan pada repository; detail diskretisasi dan engineering tidak diklaim identik dengan implementasi parent paper.\n",
        hard35,
        flags=re.S,
    )
    method = replace_section(method, "## 3.5 Preprocessing Frekuensi-Angular AF2", "## 3.6 Analisis dan Optimasi AF2", hard35)

    assembled = f"""% {TITLE}
% NAMA MAHASISWA — NIM/PROGRAM STUDI
% SEKOLAH PASCASARJANA UNIVERSITAS SUMATERA UTARA — 2026

# BAB I PENDAHULUAN

{background}

{problem}

\\newpage

# BAB II TINJAUAN PUSTAKA

{lit}

\\newpage

# BAB III METODE PENELITIAN

{method}
"""

    assembled = strip_admin(assembled)
    assembled, used = replace_citations(assembled)

    refs = [BIB[k] for k in sorted(used) if k in BIB]
    if refs:
        assembled += "\n\n\\newpage\n\n# DAFTAR PUSTAKA\n\n"
        for ref in sorted(refs, key=lambda x: x.lower()):
            assembled += ref + "\n\n"

    return assembled, used


def set_run_font(run, size: Pt | None = None, bold: bool | None = None):
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    set_run_font(run, Pt(12))


def postprocess_docx(path: Path):
    doc = Document(path)
    for section in doc.sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(3)
        section.left_margin = Cm(4)
        section.right_margin = Cm(3)
        section.header_distance = Cm(1.25)
        section.footer_distance = Cm(1.25)
        # Working release: consistent Arabic page number at bottom-right.
        footer_p = section.footer.paragraphs[0]
        footer_p.clear()
        add_page_number(footer_p)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.first_line_indent = Cm(1.27)
    normal.paragraph_format.space_after = Pt(0)

    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.first_line_indent = Cm(0)

    for p in doc.paragraphs:
        txt = p.text.strip()
        if p.style.name == "Normal":
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if p.style.name == "Heading 1" and txt.startswith("BAB "):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.page_break_before = True
        elif p.style.name in {"Heading 2", "Heading 3"}:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if txt == TITLE:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            set_run_font(r, Pt(12))

    for table in doc.tables:
        table.autofit = True
        for row_i, row in enumerate(table.rows):
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.first_line_indent = Cm(0)
                    p.paragraph_format.line_spacing = 1.0
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for r in p.runs:
                        set_run_font(r, Pt(9), bold=True if row_i == 0 else None)

    # The first three metadata paragraphs are centered and compact.
    for p in doc.paragraphs[:6]:
        if p.text.strip():
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = 1.0
            for r in p.runs:
                set_run_font(r, Pt(12), bold=True if p.text.strip() == TITLE else None)

    doc.save(path)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--repo", type=Path, default=Path("."))
    ap.add_argument("--out", type=Path, required=True)
    ap.add_argument("--emit-markdown", type=Path)
    args = ap.parse_args()

    md, used = prepare_sources(args.repo)
    if args.emit_markdown:
        args.emit_markdown.parent.mkdir(parents=True, exist_ok=True)
        args.emit_markdown.write_text(md, encoding="utf-8")

    args.out.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory() as td:
        tmp = Path(td)
        src = tmp / "proposal.md"
        src.write_text(md, encoding="utf-8")
        cmd = [
            "pandoc", str(src),
            "--from=markdown+tex_math_single_backslash+raw_tex",
            "--to=docx",
            "--standalone",
            "--toc",
            "--toc-depth=3",
            "--metadata", "lang=id-ID",
            "-o", str(args.out),
        ]
        subprocess.run(cmd, check=True)
    postprocess_docx(args.out)
    print(f"built: {args.out}")
    print(f"citation keys resolved: {len(used)}")


if __name__ == "__main__":
    main()
