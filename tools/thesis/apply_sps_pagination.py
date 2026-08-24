#!/usr/bin/env python3
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def clear_story(story):
    for p in story.paragraphs:
        p.clear()
    for table in list(story.tables):
        table._element.getparent().remove(table._element)


def set_run_font(run):
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Times New Roman")
    rpr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rpr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_page_field(paragraph, alignment):
    paragraph.clear()
    paragraph.alignment = alignment
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, display, end])
    set_run_font(run)


def set_page_number_format(section, fmt: str, start: int | None = None):
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:fmt"), fmt)
    start_key = qn("w:start")
    if start is None:
        if start_key in pg_num.attrib:
            del pg_num.attrib[start_key]
    else:
        pg_num.set(start_key, str(start))


def previous_paragraph_has_section_break(paragraph) -> bool:
    prev = paragraph._p.getprevious()
    if prev is None or prev.tag != qn("w:p"):
        return False
    ppr = prev.find(qn("w:pPr"))
    return ppr is not None and ppr.find(qn("w:sectPr")) is not None


def insert_next_page_section_break_before(paragraph):
    if previous_paragraph_has_section_break(paragraph):
        return
    new_p = OxmlElement("w:p")
    ppr = OxmlElement("w:pPr")
    sect_pr = OxmlElement("w:sectPr")
    sect_type = OxmlElement("w:type")
    sect_type.set(qn("w:val"), "nextPage")
    sect_pr.append(sect_type)
    ppr.append(sect_pr)
    new_p.append(ppr)
    paragraph._p.addprevious(new_p)


def apply_page_geometry(section):
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)
    # Official guide positions page numbers 1.5 cm from the top/bottom edge.
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)


def configure_front_matter(section):
    section.different_first_page_header_footer = False
    set_page_number_format(section, "lowerRoman", start=1)
    clear_story(section.header)
    clear_story(section.footer)
    clear_story(section.first_page_header)
    clear_story(section.first_page_footer)
    add_page_field(section.footer.paragraphs[0], WD_ALIGN_PARAGRAPH.CENTER)


def configure_body_section(section, restart_at_one: bool):
    section.different_first_page_header_footer = True
    set_page_number_format(section, "decimal", start=1 if restart_at_one else None)

    clear_story(section.header)
    clear_story(section.footer)
    clear_story(section.first_page_header)
    clear_story(section.first_page_footer)

    # Normal body pages: Arabic page number top-right.
    add_page_field(section.header.paragraphs[0], WD_ALIGN_PARAGRAPH.RIGHT)
    # Chapter-opening page: Arabic page number bottom-right.
    add_page_field(section.first_page_footer.paragraphs[0], WD_ALIGN_PARAGRAPH.RIGHT)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--docx", type=Path, required=True)
    args = ap.parse_args()

    doc = Document(args.docx)
    chapter_headings = [
        p for p in doc.paragraphs
        if p.style.name == "Heading 1" and p.text.strip().startswith("BAB ")
    ]
    if not chapter_headings:
        raise RuntimeError("no BAB Heading 1 paragraphs found; cannot sectionize pagination")

    # A next-page section break before each chapter produces:
    # section 0 = front matter, section 1 = Bab I, section 2 = Bab II, ...
    for p in chapter_headings:
        p.paragraph_format.page_break_before = False
        insert_next_page_section_break_before(p)

    temp = args.docx.with_name(args.docx.stem + ".sectionizing.tmp.docx")
    doc.save(temp)
    doc = Document(temp)

    expected_sections = len(chapter_headings) + 1
    if len(doc.sections) != expected_sections:
        raise RuntimeError(
            f"unexpected section count after pagination split: {len(doc.sections)}; "
            f"expected {expected_sections}"
        )

    for i, section in enumerate(doc.sections):
        apply_page_geometry(section)
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        section.first_page_header.is_linked_to_previous = False
        section.first_page_footer.is_linked_to_previous = False
        if i == 0:
            configure_front_matter(section)
        else:
            configure_body_section(section, restart_at_one=(i == 1))

    doc.save(args.docx)
    temp.unlink(missing_ok=True)
    print(
        f"applied SPs pagination: 1 Roman front-matter section + "
        f"{len(chapter_headings)} Arabic chapter sections"
    )


if __name__ == "__main__":
    main()
