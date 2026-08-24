#!/usr/bin/env python3
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.table import Table

import apply_sps_pagination as pag


TARGET_CAPTION_PREFIX = "Tabel 2.1"
TARGET_WIDTHS_CM = (0.9, 3.2, 2.0, 4.3, 4.0, 8.3)  # total 22.7 cm


def previous_paragraph_has_section_break(paragraph) -> bool:
    prev = paragraph._p.getprevious()
    if prev is None or prev.tag != qn("w:p"):
        return False
    ppr = prev.find(qn("w:pPr"))
    return ppr is not None and ppr.find(qn("w:sectPr")) is not None


def insert_next_page_break_before(paragraph):
    if previous_paragraph_has_section_break(paragraph):
        return
    new_p = OxmlElement("w:p")
    ppr = OxmlElement("w:pPr")
    sect_pr = OxmlElement("w:sectPr")
    sect_type = OxmlElement("w:type")
    sect_type.set(qn("w:val"), "nextPage")
    sect_pr.append(sect_type)
    ppr.append(sect_pr)
    new_p.append(ppr)
    paragraph._p.addprevious(new_p)


def insert_next_page_break_after_table(table: Table):
    nxt = table._tbl.getnext()
    if nxt is not None and nxt.tag == qn("w:p"):
        ppr = nxt.find(qn("w:pPr"))
        if ppr is not None and ppr.find(qn("w:sectPr")) is not None:
            return
    new_p = OxmlElement("w:p")
    ppr = OxmlElement("w:pPr")
    sect_pr = OxmlElement("w:sectPr")
    sect_type = OxmlElement("w:type")
    sect_type.set(qn("w:val"), "nextPage")
    sect_pr.append(sect_type)
    ppr.append(sect_pr)
    new_p.append(ppr)
    table._tbl.addnext(new_p)


def find_table_after_caption(doc: Document, caption_paragraph) -> Table:
    el = caption_paragraph._p.getnext()
    while el is not None:
        if el.tag == qn("w:tbl"):
            return Table(el, doc._body)
        if el.tag == qn("w:p"):
            ppr = el.find(qn("w:pPr"))
            if ppr is not None and ppr.find(qn("w:sectPr")) is not None:
                raise RuntimeError("section break encountered before target table")
        el = el.getnext()
    raise RuntimeError("no table found after Tabel 2.1 caption")


def section_index_for_element(doc: Document, target_element) -> int:
    idx = 0
    for child in doc._body._element:
        if child is target_element:
            return idx
        if child.tag == qn("w:p"):
            ppr = child.find(qn("w:pPr"))
            if ppr is not None and ppr.find(qn("w:sectPr")) is not None:
                idx += 1
    return idx


def configure_continuation_header(section):
    section.different_first_page_header_footer = False
    pag.set_page_number_format(section, "decimal", start=None)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.first_page_header.is_linked_to_previous = False
    section.first_page_footer.is_linked_to_previous = False
    pag.clear_story(section.header)
    pag.clear_story(section.footer)
    pag.clear_story(section.first_page_header)
    pag.clear_story(section.first_page_footer)
    pag.add_page_field(section.header.paragraphs[0], WD_ALIGN_PARAGRAPH.RIGHT)


def set_portrait_geometry(section):
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)


def set_landscape_geometry(section):
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)


def set_cell_width(cell, width_cm: float):
    width = Cm(width_cm)
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.twips)))
    tc_w.set(qn("w:type"), "dxa")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    hdr = tr_pr.find(qn("w:tblHeader"))
    if hdr is None:
        hdr = OxmlElement("w:tblHeader")
        tr_pr.append(hdr)
    hdr.set(qn("w:val"), "true")


def style_target_table(table: Table):
    if len(table.columns) != len(TARGET_WIDTHS_CM):
        raise RuntimeError(
            f"Tabel 2.1 expected {len(TARGET_WIDTHS_CM)} columns, found {len(table.columns)}"
        )
    table.autofit = False
    repeat_header(table.rows[0])
    for row_i, row in enumerate(table.rows):
        for col_i, cell in enumerate(row.cells):
            set_cell_width(cell, TARGET_WIDTHS_CM[col_i])
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(9.5)
                    rpr = run._element.get_or_add_rPr()
                    rpr.rFonts.set(qn("w:ascii"), "Times New Roman")
                    rpr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
                    rpr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    if row_i == 0:
                        run.bold = True


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--docx", type=Path, required=True)
    args = ap.parse_args()

    doc = Document(args.docx)
    captions = [p for p in doc.paragraphs if p.text.strip().startswith(TARGET_CAPTION_PREFIX)]
    if len(captions) != 1:
        raise RuntimeError(f"expected exactly one {TARGET_CAPTION_PREFIX} caption, found {len(captions)}")
    caption = captions[0]
    table = find_table_after_caption(doc, caption)

    insert_next_page_break_before(caption)
    insert_next_page_break_after_table(table)
    temp = args.docx.with_name(args.docx.stem + ".landscape.tmp.docx")
    doc.save(temp)

    doc = Document(temp)
    caption = next(p for p in doc.paragraphs if p.text.strip().startswith(TARGET_CAPTION_PREFIX))
    table = find_table_after_caption(doc, caption)
    landscape_idx = section_index_for_element(doc, caption._p)
    if landscape_idx <= 0 or landscape_idx >= len(doc.sections):
        raise RuntimeError(f"invalid landscape section index {landscape_idx}")

    set_landscape_geometry(doc.sections[landscape_idx])
    configure_continuation_header(doc.sections[landscape_idx])

    # The inserted break after the table creates a portrait continuation section
    # for the remainder of Bab II before the already-existing Bab III section.
    if landscape_idx + 1 < len(doc.sections):
        set_portrait_geometry(doc.sections[landscape_idx + 1])
        configure_continuation_header(doc.sections[landscape_idx + 1])

    caption.paragraph_format.keep_with_next = True
    caption.paragraph_format.page_break_before = False
    style_target_table(table)

    doc.save(args.docx)
    temp.unlink(missing_ok=True)
    print(
        f"isolated {TARGET_CAPTION_PREFIX} in landscape section {landscape_idx}; "
        f"table columns normalized to {sum(TARGET_WIDTHS_CM):.1f} cm"
    )


if __name__ == "__main__":
    main()
