#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

import build_proposal_docx as base


# Primary-PDF corrections / additions applied before base.prepare_sources().
# These deliberately override older ad-hoc builder metadata.
base.BIB["COF-10"] = (
    "de Oliveira, E. M., Leme, D. S., Barbosa, B. H. G., Rodarte, M. P., "
    "& Pereira, R. G. F. A. (2016). A computer vision system for coffee beans "
    "classification based on computational intelligence techniques. Journal of "
    "Food Engineering, 171, 22–27. https://doi.org/10.1016/j.jfoodeng.2015.10.009"
)
base.BIB["COF-17"] = (
    "García, M., Candelo-Becerra, J. E., & Hoyos, F. E. (2019). Quality and defect "
    "inspection of green coffee beans using a computer vision system. Applied "
    "Sciences, 9(19), 4195. https://doi.org/10.3390/app9194195"
)


def merge_adjacent_author_year_citations(text: str) -> str:
    """Merge adjacent converted citation groups such as `(A, 2024)(B, 2025)`.

    The base builder converts internal citation keys before this pass. Restrict the
    merge to parenthetical groups containing a four-digit year so ordinary math or
    explanatory parentheses are not broadly collapsed.
    """

    pattern = re.compile(
        r"\(([^()\n]*?(?:19|20)\d{2}[^()\n]*?)\)\s*"
        r"\(([^()\n]*?(?:19|20)\d{2}[^()\n]*?)\)"
    )
    prev = None
    while prev != text:
        prev = text
        text = pattern.sub(lambda m: f"({m.group(1).strip()}; {m.group(2).strip()})", text)
    return text


def normalize_table_caption(raw: str) -> str:
    """Normalize `Tabel 2.1. Judul` -> `Tabel 2.1  Judul` with no terminal period."""

    raw = raw.strip().rstrip(".")
    m = re.match(r"^(Tabel\s+\d+\.\d+)\.?\s*(.*)$", raw)
    if not m:
        return raw
    number, title = m.groups()
    return f"{number}  {title.strip()}" if title.strip() else number


def sanitize_markdown(text: str) -> str:
    # Formal manuscript does not use horizontal-rule separators from the working MD.
    text = re.sub(r"(?m)^\s*---\s*$\n?", "", text)

    # Remove fenced-code delimiters so explanatory text does not become monospaced.
    text = re.sub(r"(?m)^```[^\n]*\n?", "", text)
    text = re.sub(r"(?m)^```\s*$\n?", "", text)

    # Table titles must be captions, not Heading paragraphs/TOC entries. The SPs
    # convention also omits a period after the table number and at line end.
    text = re.sub(
        r"(?m)^#{2,6}\s+(Tabel\s+\d+\.\d+\.?\s+.+?)\s*$",
        lambda m: f"**{normalize_table_caption(m.group(1))}**",
        text,
    )

    # SPs guide disallows bullet-symbol lists in the formal manuscript. Repeated
    # `1.` markers are intentionally used; Pandoc numbers them automatically.
    text = re.sub(r"(?m)^(\s*)[-*]\s+", r"\g<1>1. ", text)

    text = merge_adjacent_author_year_citations(text)
    return text


def set_run_font(run, size: Pt):
    run.font.name = "Times New Roman"
    run.font.size = size
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Times New Roman")
    rpr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rpr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def find_heading_style(doc: Document, level: int):
    """Resolve a Word heading by stable styleId before falling back to display name.

    Pandoc/Word can localize the visible style name (for example `Judul 1`) while
    preserving the built-in styleId `Heading1`. The first-generation builder assumed
    the English display name and therefore failed on GitHub Actions. This resolver
    treats styleId as the stable identity and keeps the visible name only as fallback.
    """

    target_id = f"heading{level}"
    target_name = f"heading {level}"
    for style in doc.styles:
        style_id = (getattr(style, "style_id", "") or "").casefold().replace(" ", "")
        if style_id == target_id:
            return style
    for style in doc.styles:
        name = (getattr(style, "name", "") or "").casefold().strip()
        if name == target_name:
            return style
    return None


def normalize_heading_style_names(path: Path):
    """Normalize Heading1/2/3 display names before legacy post-processing.

    Existing heading paragraphs remain attached to the same style objects; only the
    localized display names are normalized so legacy code and subsequent passes that
    compare `paragraph.style.name` continue to work consistently.
    """

    doc = Document(path)
    resolved = []
    for level in (1, 2, 3):
        style = find_heading_style(doc, level)
        if style is None:
            available = ", ".join(
                f"{getattr(s, 'style_id', '?')}={getattr(s, 'name', '?')}"
                for s in doc.styles
                if "heading" in (getattr(s, "style_id", "") or "").casefold()
                or "judul" in (getattr(s, "name", "") or "").casefold()
            )
            raise RuntimeError(
                f"cannot resolve built-in heading level {level}; candidates: {available or 'none'}"
            )
        expected = f"Heading {level}"
        if style.name != expected:
            style.name = expected
        resolved.append(f"{style.style_id}->{style.name}")
    doc.save(path)
    print("normalized Word heading styles: " + ", ".join(resolved))


def split_chapter_heading(paragraph):
    text = paragraph.text.strip()
    m = re.fullmatch(r"BAB\s+([IVXLCDM]+)\s+(.+)", text)
    if not m:
        return False
    paragraph.clear()
    r1 = paragraph.add_run(f"BAB {m.group(1)}")
    r1.add_break()
    paragraph.add_run(m.group(2).upper())
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.page_break_before = True
    for r in paragraph.runs:
        set_run_font(r, Pt(12))
        r.bold = True
    return True


def formal_postprocess(path: Path):
    """Low-risk formatting corrections before rendered-layout QA."""

    doc = Document(path)

    for p in doc.paragraphs:
        text = p.text.strip()
        if p.style.name == "Heading 1":
            split_chapter_heading(p)

        if text.startswith("Tabel "):
            # Table captions are ordinary centered single-spaced paragraphs.
            try:
                p.style = doc.styles["Normal"]
            except KeyError:
                pass
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.keep_with_next = True
            for r in p.runs:
                set_run_font(r, Pt(12))
                r.bold = True

    for table in doc.tables:
        wide = len(table.columns) >= 6
        font_size = Pt(9.5 if wide else 11)
        for row_i, row in enumerate(table.rows):
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.first_line_indent = Cm(0)
                    p.paragraph_format.line_spacing = 1.0
                    p.paragraph_format.space_after = Pt(0)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for r in p.runs:
                        set_run_font(r, font_size)
                        if row_i == 0:
                            r.bold = True

    doc.save(path)


def validate_source(md: str, used: set[str]):
    missing_bib = sorted(k for k in used if k not in base.BIB)
    if missing_bib:
        raise RuntimeError(
            "bibliography records missing for used citation keys: " + ", ".join(missing_bib)
        )

    unresolved = sorted(
        set(
            re.findall(
                r"\[(?:STD|REV|COF(?:-SUP)?|DET|EVAL|DIAG|FG|PRE|SPEC|FREQ|WAVE|THEORY)-\d+[^\]]*\]",
                md,
            )
        )
    )
    if unresolved:
        raise RuntimeError(
            "unresolved internal citation tokens remain after assembly: "
            + "; ".join(unresolved[:10])
        )

    if len(used) < 50:
        print(
            f"WARNING: assembled proposal currently cites {len(used)} unique keyed sources; "
            "the final SPs reference-count requirement is not yet treated as closed. "
            "Do not pad references artificially."
        )

    if re.search(r"(?m)^\s*[-*]\s+", md):
        raise RuntimeError("unordered Markdown bullet marker remains after sanitization")

    if re.search(r"(?m)^```", md):
        raise RuntimeError("fenced-code marker remains after sanitization")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--repo", type=Path, default=Path("."))
    ap.add_argument("--out", type=Path, required=True)
    ap.add_argument("--emit-markdown", type=Path)
    ap.add_argument(
        "--reference-doc",
        type=Path,
        help="Optional campus DOCX scaffold/reference document for Pandoc. Formal local releases should supply the uploaded template.",
    )
    args = ap.parse_args()

    md, used = base.prepare_sources(args.repo)
    md = sanitize_markdown(md)
    validate_source(md, used)

    if args.emit_markdown:
        args.emit_markdown.parent.mkdir(parents=True, exist_ok=True)
        args.emit_markdown.write_text(md, encoding="utf-8")

    args.out.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory() as td:
        src = Path(td) / "proposal.md"
        src.write_text(md, encoding="utf-8")
        cmd = [
            "pandoc",
            str(src),
            "--from=markdown+tex_math_single_backslash+raw_tex",
            "--to=docx",
            "--standalone",
            "--toc",
            "--toc-depth=3",
            "--metadata",
            "lang=id-ID",
        ]
        if args.reference_doc:
            if not args.reference_doc.exists():
                raise FileNotFoundError(args.reference_doc)
            cmd.extend(["--reference-doc", str(args.reference_doc)])
        else:
            print(
                "WARNING: no --reference-doc supplied; this is a working build, not a template-faithful formal release."
            )
        cmd.extend(["-o", str(args.out)])
        subprocess.run(cmd, check=True)

    # Pandoc may localize the visible names of built-in Word heading styles on CI.
    # Normalize them before invoking the first-generation postprocessor, which still
    # uses English display names internally.
    normalize_heading_style_names(args.out)

    # Retain the original margin/body styling pass, then apply QA-hardened corrections.
    base.postprocess_docx(args.out)
    formal_postprocess(args.out)

    print(f"built: {args.out}")
    print(f"citation keys resolved: {len(used)}")
    print(
        "OPEN LAYOUT GATES: template-faithful build, actual rendered-page QA, "
        "narrative citations, canonical APA bibliography"
    )


if __name__ == "__main__":
    main()
