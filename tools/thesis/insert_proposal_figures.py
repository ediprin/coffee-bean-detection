#!/usr/bin/env python3
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

# SPs guide convention used here:
# - caption below figure;
# - centered;
# - no period after the figure number;
# - no terminal period;
# - title case;
# - TNR 12 unless a verified exception is introduced later.
FIGURES = {
    "Gambar 3.1 pada dokumen final:": (
        "fig_3_1_research_framework.png",
        "Gambar 3.1  Kerangka Penelitian",
    ),
    "Gambar 3.2 pada dokumen final:": (
        "fig_3_2_native_vs_af2.png",
        "Gambar 3.2  Arsitektur Native YOLO26 dan AF2–YOLO26",
    ),
    "Gambar 3.3 pada dokumen final:": (
        "fig_3_3_af2_operator.png",
        "Gambar 3.3  Alur Preprocessing Frekuensi-Angular AF2",
    ),
}


def clear_paragraph(p):
    p_el = p._element
    for child in list(p_el):
        if child.tag != qn("w:pPr"):
            p_el.remove(child)


def insert_paragraph_after(paragraph, text=""):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph

    p = Paragraph(new_p, paragraph._parent)
    if text:
        p.add_run(text)
    return p


def style_caption(p, text: str):
    p.text = text
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    for r in p.runs:
        r.font.name = "Times New Roman"
        r._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
        r._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
        r.font.size = Pt(12)


def add_picture_in_paragraph(p, image: Path, width_cm: float):
    clear_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(image), width=Cm(width_cm))


def insert_fig_34(doc: Document, image_dir: Path):
    image = image_dir / "fig_3_4_optimization_genealogy.png"
    if not image.exists():
        return False
    for p in doc.paragraphs:
        if p.text.strip().startswith("3.6 Analisis dan Optimasi AF2"):
            img_p = insert_paragraph_after(p)
            add_picture_in_paragraph(img_p, image, 13.5)
            cap = insert_paragraph_after(img_p)
            style_caption(
                cap,
                "Gambar 3.4  Genealogi Optimasi AF2 dan Eksperimen Konfirmatori",
            )
            return True
    return False


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--docx", type=Path, required=True)
    ap.add_argument("--image-dir", type=Path, required=True)
    args = ap.parse_args()

    doc = Document(args.docx)
    found = {k: False for k in FIGURES}

    for p in doc.paragraphs:
        t = p.text.strip()
        for marker, (filename, caption) in FIGURES.items():
            if t.startswith(marker):
                image = args.image_dir / filename
                if not image.exists():
                    raise FileNotFoundError(image)
                add_picture_in_paragraph(p, image, 13.5)
                cap = insert_paragraph_after(p)
                style_caption(cap, caption)
                found[marker] = True
                break

    found_34 = insert_fig_34(doc, args.image_dir)
    missing = [k for k, v in found.items() if not v]
    if missing:
        raise RuntimeError(f"figure placeholders not found: {missing}")
    if not found_34:
        raise RuntimeError("heading 3.6 not found for Figure 3.4 insertion")

    doc.save(args.docx)
    print(f"inserted Figures 3.1–3.4 into {args.docx}")


if __name__ == "__main__":
    main()
