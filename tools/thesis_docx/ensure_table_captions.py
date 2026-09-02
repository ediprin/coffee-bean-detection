#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

FONT = "Times New Roman"

# Identify formal thesis tables by their header row. The pruned proposal keeps
# one canonical table for each distinct purpose and does not repeat B0-B3 in
# the seed-confirmation subsection.
CAPTIONS = {
    ("No.", "Penulis dan Tahun", "Sumber Publikasi/Venue", "Fokus Penelitian", "Metode/Model", "Relevansi dengan Penelitian"):
        "Tabel 2.1: Penelitian Terkait",
    ("Dataset", "Sumber", "Peran"):
        "Tabel 3.1: Sumber dan Peran Dataset Penelitian",
    ("Dataset", "Versi", "Task", "Jumlah kelas"):
        "Tabel 3.2: Dataset Publik untuk Konfirmasi",
    ("Kode", "Kondisi", "Peran dalam eksperimen"):
        "Tabel 3.3: Kondisi Utama Eksperimen",
    ("Konfigurasi", "Perubahan utama", "Tujuan pengujian"):
        "Tabel 3.4: Variasi Desain Prapemrosesan",
    ("Parameter", "Nilai"):
        "Tabel 3.5: Konfigurasi Utama Pelatihan YOLO26n",
}

TEXT_REPLACEMENTS = {
    "Konfigurasi utama YOLO26n ditunjukkan pada Tabel 3.3.":
        "Konfigurasi utama YOLO26n ditunjukkan pada Tabel 3.5.",
}


def set_run_font(run, *, bold=None) -> None:
    run.font.name = FONT
    run.font.size = Pt(12)
    if bold is not None:
        run.bold = bold
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), FONT)
    for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        key = qn(f"w:{attr}")
        if key in rfonts.attrib:
            del rfonts.attrib[key]


def is_equation_table(table) -> bool:
    has_math = bool(
        list(table._tbl.iter(qn("m:oMath")))
        or list(table._tbl.iter(qn("m:oMathPara")))
    )
    return (
        len(table.rows) == 1
        and len(table.columns) == 3
        and has_math
        and bool(re.fullmatch(r"\(\d+(?:\.\d+)?\)", table.cell(0, 2).text.strip()))
    )


def header_key(table) -> tuple[str, ...]:
    return tuple(cell.text.strip() for cell in table.rows[0].cells)


def set_caption_format(paragraph, caption_style) -> None:
    paragraph.style = caption_style
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = paragraph.paragraph_format
    pf.first_line_indent = None
    pf.left_indent = None
    pf.right_indent = None
    pf.line_spacing = 1.0
    pf.space_before = Pt(6)
    pf.space_after = Pt(3)
    for run in paragraph.runs:
        set_run_font(run, bold=True)


def previous_paragraph_node(table_node):
    node = table_node.getprevious()
    while node is not None:
        if node.tag == qn("w:p"):
            return node
        node = node.getprevious()
    return None


def paragraph_text(node) -> str:
    return "".join(node.xpath(".//w:t/text()")).strip()


def insert_caption_before(doc, table, caption_text: str, caption_style):
    prev_node = previous_paragraph_node(table._tbl)
    prev_text = paragraph_text(prev_node) if prev_node is not None else ""

    if re.match(r"^Tabel\s+\d+\.\d+", prev_text, re.I):
        paragraph = next(p for p in doc.paragraphs if p._p is prev_node)
        paragraph.text = caption_text
        set_caption_format(paragraph, caption_style)
        return paragraph

    paragraph = doc.add_paragraph()
    paragraph.text = caption_text
    set_caption_format(paragraph, caption_style)
    table._tbl.addprevious(paragraph._p)
    return paragraph


def apply(input_path: Path, output_path: Path) -> None:
    doc = Document(input_path)
    if "Caption Table" not in doc.styles:
        raise RuntimeError("Caption Table style is missing")
    caption_style = doc.styles["Caption Table"]

    for paragraph in doc.paragraphs:
        replacement = TEXT_REPLACEMENTS.get(paragraph.text.strip())
        if replacement:
            paragraph.text = replacement
            for run in paragraph.runs:
                set_run_font(run)

    found = []
    unknown = []
    for table in doc.tables:
        if is_equation_table(table):
            continue
        key = header_key(table)
        caption = CAPTIONS.get(key)
        if caption is None:
            unknown.append(key)
            continue
        insert_caption_before(doc, table, caption, caption_style)
        found.append(caption)

    if unknown:
        raise RuntimeError(f"Unmapped regular thesis table headers: {unknown}")
    if len(found) != len(CAPTIONS):
        raise RuntimeError(f"Expected {len(CAPTIONS)} regular thesis tables, found {len(found)}: {found}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print("Formal table captions ensured:")
    for caption in found:
        print(f"- {caption}")


def main() -> None:
    parser = argparse.ArgumentParser(description="Ensure all formal thesis tables have deterministic numbered captions.")
    parser.add_argument("input")
    parser.add_argument("--output", default=None)
    args = parser.parse_args()
    source = Path(args.input)
    target = Path(args.output) if args.output else source
    apply(source, target)


if __name__ == "__main__":
    main()
