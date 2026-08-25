#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

# Flowchart penelitian berbentuk portrait sehingga lebih proporsional pada 8.5 cm.
# Diagram arsitektur dua-panel membutuhkan lebar lebih besar agar label tetap terbaca.
PORTRAIT_MAX_WIDTH = Cm(8.5)
LANDSCAPE_MAX_WIDTH = Cm(12.5)
MAX_FIGURE_HEIGHT = Cm(13.5)


def _paragraph_has_drawing(paragraph) -> bool:
    return bool(
        list(paragraph._p.iter(qn("w:drawing")))
        or list(paragraph._p.iter(qn("w:pict")))
    )


def _remove_duplicate_pandoc_image_captions(doc) -> int:
    """Remove Pandoc alt-text captions when a numbered thesis caption follows."""
    paragraphs = list(doc.paragraphs)
    removed = 0
    for index, paragraph in enumerate(paragraphs):
        if paragraph.style.name != "Image Caption":
            continue
        if index == 0 or index + 1 >= len(paragraphs):
            continue
        previous = paragraphs[index - 1]
        following = paragraphs[index + 1]
        if not _paragraph_has_drawing(previous):
            continue
        if following.style.name != "Caption Figure":
            continue
        if not re.match(r"^Gambar\s+\d+\.\d+\.", following.text.strip(), re.I):
            continue
        paragraph._element.getparent().remove(paragraph._element)
        removed += 1
    return removed


def format_figures(input_path: Path, output_path: Path) -> None:
    doc = Document(input_path)

    figure_paragraphs = 0
    for paragraph in doc.paragraphs:
        if not _paragraph_has_drawing(paragraph):
            continue
        figure_paragraphs += 1
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.left_indent = Cm(0)
        paragraph.paragraph_format.right_indent = Cm(0)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)

    resized = 0
    for shape in doc.inline_shapes:
        max_width = PORTRAIT_MAX_WIDTH if shape.height > shape.width else LANDSCAPE_MAX_WIDTH
        width_scale = max_width / shape.width if shape.width > max_width else 1.0
        height_scale = MAX_FIGURE_HEIGHT / shape.height if shape.height > MAX_FIGURE_HEIGHT else 1.0
        scale = min(width_scale, height_scale)
        if scale >= 1.0:
            continue
        shape.width = int(shape.width * scale)
        shape.height = int(shape.height * scale)
        resized += 1

    removed_captions = _remove_duplicate_pandoc_image_captions(doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(
        f"Formatted {figure_paragraphs} image paragraphs; "
        f"resized {resized} figures; "
        f"removed {removed_captions} duplicate Pandoc image captions."
    )


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Center and constrain Markdown-generated figures in thesis DOCX."
    )
    parser.add_argument("input")
    parser.add_argument("--output", default=None)
    args = parser.parse_args()
    source = Path(args.input)
    target = Path(args.output) if args.output else source
    format_figures(source, target)


if __name__ == "__main__":
    main()
