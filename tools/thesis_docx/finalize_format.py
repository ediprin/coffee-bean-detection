#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from ensure_table_captions import apply as ensure_table_captions

FONT_NAME = "Times New Roman"
TABLE_EDGES = ("top", "left", "bottom", "right", "insideH", "insideV")
CELL_EDGES = ("top", "left", "bottom", "right")


def _replace_borders(parent, tag_name: str, edges, visible: bool) -> None:
    old = parent.find(qn(tag_name))
    if old is not None:
        parent.remove(old)
    borders = OxmlElement(tag_name)
    for edge in edges:
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single" if visible else "nil")
        if visible:
            element.set(qn("w:sz"), "8")
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), "000000")
        borders.append(element)
    parent.append(borders)


def _set_table_style(table, style_name: str | None) -> None:
    """Set an explicit Word table style without depending on the document theme."""
    tbl_pr = table._tbl.tblPr
    tbl_style = tbl_pr.find(qn("w:tblStyle"))
    if style_name is None:
        if tbl_style is not None:
            tbl_pr.remove(tbl_style)
        return
    if tbl_style is None:
        tbl_style = OxmlElement("w:tblStyle")
        tbl_pr.insert(0, tbl_style)
    tbl_style.set(qn("w:val"), style_name)


def _set_table_borders(table, visible: bool) -> None:
    tbl_pr = table._tbl.tblPr
    _set_table_style(table, "TableGrid" if visible else None)
    _replace_borders(tbl_pr, "w:tblBorders", TABLE_EDGES, visible)
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            _replace_borders(tc_pr, "w:tcBorders", CELL_EDGES, visible)


def _is_equation_layout_table(table) -> bool:
    """Identify only the synthetic 1x3 table used to center and number equations."""
    if len(table.rows) != 1 or len(table.columns) != 3:
        return False
    has_math = bool(
        list(table._tbl.iter(qn("m:oMath")))
        or list(table._tbl.iter(qn("m:oMathPara")))
    )
    if not has_math:
        return False
    right_text = table.cell(0, 2).text.strip()
    return bool(re.fullmatch(r"\(\d+(?:\.\d+)?\)", right_text))


def _ensure_rfonts(r_pr) -> None:
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), FONT_NAME)
    for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        key = qn(f"w:{attr}")
        if key in r_fonts.attrib:
            del r_fonts.attrib[key]


def _force_font_in_root(root) -> None:
    for run in root.iter(qn("w:r")):
        r_pr = run.find(qn("w:rPr"))
        if r_pr is None:
            r_pr = OxmlElement("w:rPr")
            run.insert(0, r_pr)
        _ensure_rfonts(r_pr)

    for math_run in root.iter(qn("m:r")):
        r_pr = math_run.find(qn("w:rPr"))
        if r_pr is None:
            r_pr = OxmlElement("w:rPr")
            math_run.insert(0, r_pr)
        _ensure_rfonts(r_pr)

    for r_pr in root.iter(qn("w:rPr")):
        _ensure_rfonts(r_pr)


def _force_style_fonts(doc) -> None:
    styles_root = doc.styles.element
    doc_defaults = styles_root.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles_root.insert(0, doc_defaults)
    r_pr_default = doc_defaults.find(qn("w:rPrDefault"))
    if r_pr_default is None:
        r_pr_default = OxmlElement("w:rPrDefault")
        doc_defaults.append(r_pr_default)
    r_pr = r_pr_default.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        r_pr_default.append(r_pr)
    _ensure_rfonts(r_pr)

    for style in styles_root.findall(qn("w:style")):
        r_pr = style.find(qn("w:rPr"))
        if r_pr is None:
            r_pr = OxmlElement("w:rPr")
            style.append(r_pr)
        _ensure_rfonts(r_pr)


def _set_math_default_font(doc) -> None:
    settings = doc.settings._element
    math_pr = settings.find(qn("m:mathPr"))
    if math_pr is None:
        math_pr = OxmlElement("m:mathPr")
        settings.append(math_pr)
    math_font = math_pr.find(qn("m:mathFont"))
    if math_font is None:
        math_font = OxmlElement("m:mathFont")
        math_pr.insert(0, math_font)
    math_font.set(qn("m:val"), FONT_NAME)


def _format_source_code_blocks(doc) -> None:
    """Make Markdown fenced flow diagrams readable in Word."""
    for paragraph in doc.paragraphs:
        if paragraph.style.name != "Source Code":
            continue
        lines = []
        for line in paragraph.text.splitlines():
            clean = re.sub(r"[ \t]{2,}", " ", line.strip())
            if clean:
                lines.append(clean)
        paragraph.text = "\n".join(lines)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.left_indent = Cm(0)
        paragraph.paragraph_format.right_indent = Cm(0)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
        for run in paragraph.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(12)


def _restore_false_table_caption_paragraphs(doc) -> None:
    """Restore prose beginning with a table number that Pandoc post-formatting mistook for a caption."""
    for paragraph in doc.paragraphs:
        if paragraph.style.name != "Caption Table":
            continue
        match = re.match(r"^Tabel\s+(\d+\.\d+):\s+([a-z].*)$", paragraph.text.strip())
        if not match:
            continue
        paragraph.text = f"Tabel {match.group(1)} {match.group(2)}"
        paragraph.style = doc.styles["Normal"]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = paragraph.paragraph_format
        pf.first_line_indent = Cm(1.27)
        pf.left_indent = Cm(0)
        pf.right_indent = Cm(0)
        pf.line_spacing = 1.5
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(12)
            run.bold = False


def _set_cell_width(cell, width_cm: float) -> None:
    """Set a deterministic cell width in twips for Word and LibreOffice."""
    width = Cm(width_cm)
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width.twips))
    tc_w.set(qn("w:type"), "dxa")


def _prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    cant_split.set(qn("w:val"), "true")


def _format_equation_layout_table(table) -> None:
    """Lock the synthetic equation-number table to the full 14 cm text width."""
    widths_cm = (0.5, 12.0, 1.5)
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(Cm(sum(widths_cm)).twips))
    tbl_w.set(qn("w:type"), "dxa")

    grid_cols = table._tbl.tblGrid.gridCol_lst
    for index, grid_col in enumerate(grid_cols[: len(widths_cm)]):
        grid_col.set(qn("w:w"), str(Cm(widths_cm[index]).twips))

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            _set_cell_width(cell, widths_cm[index])
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = Cm(0)
                paragraph.paragraph_format.left_indent = Cm(0)
                paragraph.paragraph_format.right_indent = Cm(0)
                paragraph.paragraph_format.line_spacing = 1.0


def _format_related_research_table(doc) -> None:
    """Give Tabel 2.1 readable column proportions and keep each study row intact."""
    widths_cm = (0.75, 1.85, 2.20, 2.20, 2.00, 5.00)

    for table in doc.tables:
        if not table.rows or len(table.columns) != 6:
            continue
        headers = [cell.text.strip().replace("\n", " ") for cell in table.rows[0].cells]
        is_related_table = (
            headers[0].startswith("No")
            and "Penulis" in headers[1]
            and "Venue" in headers[2]
            and "Fokus" in headers[3]
            and "Metode" in headers[4]
            and ("Kontribusi" in headers[5] or "Relevansi" in headers[5])
        )
        if not is_related_table:
            continue

        table.autofit = False
        tbl_pr = table._tbl.tblPr

        layout = tbl_pr.find(qn("w:tblLayout"))
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            tbl_pr.append(layout)
        layout.set(qn("w:type"), "fixed")

        tbl_w = tbl_pr.find(qn("w:tblW"))
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:w"), str(Cm(sum(widths_cm)).twips))
        tbl_w.set(qn("w:type"), "dxa")

        grid_cols = table._tbl.tblGrid.gridCol_lst
        for index, grid_col in enumerate(grid_cols[: len(widths_cm)]):
            grid_col.set(qn("w:w"), str(Cm(widths_cm[index]).twips))

        for row in table.rows:
            _prevent_row_split(row)
            for column_index, cell in enumerate(row.cells):
                _set_cell_width(cell, widths_cm[column_index])
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.first_line_indent = Cm(0)
                    paragraph.paragraph_format.left_indent = Cm(0)
                    paragraph.paragraph_format.right_indent = Cm(0)
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.alignment = (
                        WD_ALIGN_PARAGRAPH.CENTER
                        if column_index in (0, 1)
                        else WD_ALIGN_PARAGRAPH.LEFT
                    )
                    for run in paragraph.runs:
                        run.font.name = FONT_NAME
                        run.font.size = Pt(9)

        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)
        break


def finalize(input_path: Path, output_path: Path) -> None:
    doc = Document(input_path)

    equation_tables = 0
    normal_tables = 0
    for table in doc.tables:
        if _is_equation_layout_table(table):
            _set_table_borders(table, visible=False)
            _format_equation_layout_table(table)
            equation_tables += 1
        else:
            _set_table_borders(table, visible=True)
            normal_tables += 1

    _format_source_code_blocks(doc)
    _restore_false_table_caption_paragraphs(doc)
    _format_related_research_table(doc)
    _force_style_fonts(doc)
    _set_math_default_font(doc)
    _force_font_in_root(doc.element)

    seen_parts = set()
    for section in doc.sections:
        for container in (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            part = container.part
            if id(part) in seen_parts:
                continue
            seen_parts.add(id(part))
            _force_font_in_root(part.element)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    ensure_table_captions(output_path, output_path)

    print(
        f"Finalized {output_path}: {normal_tables} regular tables with TableGrid/full borders, "
        f"{equation_tables} equation-layout tables without borders; font={FONT_NAME}."
    )


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Final DOCX formatting pass for SPs USU proposal builds."
    )
    parser.add_argument("input")
    parser.add_argument("--output", default=None)
    args = parser.parse_args()
    source = Path(args.input)
    target = Path(args.output) if args.output else source
    finalize(source, target)


if __name__ == "__main__":
    main()
