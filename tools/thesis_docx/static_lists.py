#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
from pathlib import Path

import fitz
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

FONT = "Times New Roman"
ENTRY_STYLE = "Static List Entry"
ANCHOR_STYLE = "Static List Anchor"


def normalize(text: str) -> str:
    text = text.replace("\u00a0", " ").replace("–", "-").replace("—", "-").replace("−", "-")
    return re.sub(r"\s+", " ", text).strip().casefold()


def pdf_pages(pdf_path: Path) -> list[str]:
    pdf = fitz.open(pdf_path)
    return [normalize(page.get_text()) for page in pdf]


def find_page(pages: list[str], text: str, start_page: int = 1) -> int:
    query = normalize(text)
    for page_number, page_text in enumerate(pages[start_page - 1 :], start=start_page):
        if query in page_text:
            return page_number
    short = query[:45]
    for page_number, page_text in enumerate(pages[start_page - 1 :], start=start_page):
        if short and short in page_text:
            return page_number
    raise RuntimeError(f"Could not locate caption in rendered PDF: {text!r}")


def find_bab1_page(pages: list[str]) -> int:
    for page_number, page_text in enumerate(pages, start=1):
        if "bab i" in page_text and "pendahuluan" in page_text and "1.1 latar belakang" in page_text:
            return page_number
    raise RuntimeError("Could not locate BAB I first page")


def ensure_style(doc, name: str, *, entry: bool):
    if name in doc.styles:
        style = doc.styles[name]
    else:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]

    style.font.name = FONT
    style.font.size = Pt(12)
    style.font.bold = False
    style.font.color.rgb = RGBColor(0, 0, 0)
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), FONT)

    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.line_spacing = 1.0
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(0)
    pf.right_indent = Cm(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(3 if entry else 0)
    try:
        pf.tab_stops.clear_all()
    except Exception:
        pass
    if entry:
        pf.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    return style


def paragraph_style_id(node) -> str | None:
    ppr = node.find(qn("w:pPr"))
    if ppr is None:
        return None
    pstyle = ppr.find(qn("w:pStyle"))
    return pstyle.get(qn("w:val")) if pstyle is not None else None


def paragraph_text(node) -> str:
    return "".join(node.xpath(".//w:t/text()")).strip()


def contains_caption_toc(node, expected_style: str) -> bool:
    return any(
        "TOC" in (item.text or "") and expected_style in (item.text or "")
        for item in node.iter(qn("w:instrText"))
    )


def clear_paragraph_content_preserve_properties(node) -> None:
    for child in list(node):
        if child.tag != qn("w:pPr"):
            node.remove(child)


def set_paragraph_style(node, style_id: str) -> None:
    ppr = node.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        node.insert(0, ppr)
    pstyle = ppr.find(qn("w:pStyle"))
    if pstyle is None:
        pstyle = OxmlElement("w:pStyle")
        ppr.insert(0, pstyle)
    pstyle.set(qn("w:val"), style_id)


def prepare_anchor_and_remove_old_entries(title_paragraph, boundary_text: str, expected_style: str, entry_style_id: str, anchor_style_id: str):
    """Keep the original placeholder paragraph as an anchor.

    The DAFTAR GAMBAR placeholder is also the paragraph that carries the section
    break immediately before BAB I. Deleting it would collapse the thesis section
    structure. We therefore clear its field content but preserve its paragraph
    properties (including w:sectPr), then insert visible list entries before it.
    """

    title_node = title_paragraph._p
    node = title_node.getnext()
    anchor = None

    while node is not None:
        if node.tag != qn("w:p"):
            node = node.getnext()
            continue
        text = paragraph_text(node).upper()
        if text == boundary_text.upper():
            break

        style_id = paragraph_style_id(node)
        next_node = node.getnext()
        if style_id == entry_style_id:
            node.getparent().remove(node)
        elif style_id == anchor_style_id:
            anchor = node
        elif contains_caption_toc(node, expected_style):
            clear_paragraph_content_preserve_properties(node)
            set_paragraph_style(node, anchor_style_id)
            anchor = node
        node = next_node

    if anchor is None:
        raise RuntimeError(
            f"Could not locate preserved list anchor after {title_paragraph.text!r}; "
            "refusing to risk deleting a section break."
        )
    return anchor


def build_entry_node(text: str, page: str, style_id: str):
    paragraph = OxmlElement("w:p")
    ppr = OxmlElement("w:pPr")
    pstyle = OxmlElement("w:pStyle")
    pstyle.set(qn("w:val"), style_id)
    ppr.append(pstyle)
    paragraph.append(ppr)

    text_run = OxmlElement("w:r")
    text_node = OxmlElement("w:t")
    text_node.text = text
    text_run.append(text_node)
    paragraph.append(text_run)

    tab_run = OxmlElement("w:r")
    tab_run.append(OxmlElement("w:tab"))
    paragraph.append(tab_run)

    page_run = OxmlElement("w:r")
    page_node = OxmlElement("w:t")
    page_node.text = page
    page_run.append(page_node)
    paragraph.append(page_run)
    return paragraph


def insert_entries_before_anchor(anchor, entries, style_id: str) -> None:
    for text, page in entries:
        anchor.addprevious(build_entry_node(text, page, style_id))


def build(input_path: Path, pdf_path: Path, output_path: Path) -> None:
    doc = Document(input_path)
    entry_style = ensure_style(doc, ENTRY_STYLE, entry=True)
    anchor_style = ensure_style(doc, ANCHOR_STYLE, entry=False)
    pages = pdf_pages(pdf_path)
    bab1_page = find_bab1_page(pages)

    table_captions = [p.text.strip() for p in doc.paragraphs if p.style.name == "Caption Table" and p.text.strip()]
    figure_captions = [p.text.strip() for p in doc.paragraphs if p.style.name == "Caption Figure" and p.text.strip()]
    if not table_captions:
        raise RuntimeError("No Caption Table paragraphs found; DAFTAR TABEL would be empty")
    if not figure_captions:
        raise RuntimeError("No Caption Figure paragraphs found; DAFTAR GAMBAR would be empty")

    table_entries = [(caption, str(find_page(pages, caption, bab1_page) - bab1_page + 1)) for caption in table_captions]
    figure_entries = [(caption, str(find_page(pages, caption, bab1_page) - bab1_page + 1)) for caption in figure_captions]

    table_title = next(p for p in doc.paragraphs if p.text.strip().upper() == "DAFTAR TABEL")
    figure_title = next(p for p in doc.paragraphs if p.text.strip().upper() == "DAFTAR GAMBAR")

    table_anchor = prepare_anchor_and_remove_old_entries(
        table_title,
        "DAFTAR GAMBAR",
        "Caption Table",
        entry_style.style_id,
        anchor_style.style_id,
    )
    figure_anchor = prepare_anchor_and_remove_old_entries(
        figure_title,
        "BAB I",
        "Caption Figure",
        entry_style.style_id,
        anchor_style.style_id,
    )
    insert_entries_before_anchor(table_anchor, table_entries, entry_style.style_id)
    insert_entries_before_anchor(figure_anchor, figure_entries, entry_style.style_id)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print("DAFTAR TABEL:")
    for text, page in table_entries:
        print(f"- {text} -> {page}")
    print("DAFTAR GAMBAR:")
    for text, page in figure_entries:
        print(f"- {text} -> {page}")


def main() -> None:
    parser = argparse.ArgumentParser(description="Materialize visible thesis lists of tables and figures.")
    parser.add_argument("input")
    parser.add_argument("pdf")
    parser.add_argument("output")
    args = parser.parse_args()
    build(Path(args.input), Path(args.pdf), Path(args.output))


if __name__ == "__main__":
    main()
