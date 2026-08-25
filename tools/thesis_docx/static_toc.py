#!/usr/bin/env python3
from __future__ import annotations

import argparse
import copy
import re
from pathlib import Path

import fitz
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

FONT = "Times New Roman"


def style_by_name(doc, name):
    return next(style for style in doc.styles if style.name == name)


def add_style(doc, name, base="Normal"):
    if name in doc.styles:
        return style_by_name(doc, name)
    style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = style_by_name(doc, base)
    return style


def set_style_font(style, bold=False):
    style.font.name = FONT
    style.font.size = Pt(12)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), FONT)


def ensure_styles(doc):
    specs = {
        "Static TOC 1": (True, 0.0, 3),
        "Static TOC 2": (False, 0.75, 1),
        "Static TOC 3": (False, 1.50, 1),
    }
    for name, (bold, indent_cm, after_pt) in specs.items():
        style = add_style(doc, name)
        set_style_font(style, bold)
        pf = style.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.line_spacing = 1.0
        pf.first_line_indent = Cm(0)
        pf.left_indent = Cm(indent_cm)
        pf.right_indent = Cm(0)
        pf.space_before = Pt(0)
        pf.space_after = Pt(after_pt)
        try:
            pf.tab_stops.clear_all()
        except Exception:
            pass
        pf.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)


def set_page_number_format(section, fmt, start=None):
    sect_pr = section._sectPr
    pg = sect_pr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType")
        sect_pr.append(pg)
    pg.set(qn("w:fmt"), fmt)
    if start is not None:
        pg.set(qn("w:start"), str(start))
    elif qn("w:start") in pg.attrib:
        del pg.attrib[qn("w:start")]


def clear_container(container):
    for paragraph in list(container.paragraphs):
        paragraph._element.getparent().remove(paragraph._element)
    for table in list(container.tables):
        table._element.getparent().remove(table._element)


def add_field(paragraph, instruction, result="1"):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    result_run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = result
    result_run.append(text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, result_run, end])


def add_page_number(paragraph, center=False):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.RIGHT
    add_field(paragraph, " PAGE ", "1")
    for run in paragraph.runs:
        run.font.name = FONT
        run.font.size = Pt(12)


def add_section_break_before(paragraph, template_sectpr):
    previous = paragraph._p.getprevious()
    if previous is None:
        return
    if previous.tag != qn("w:p"):
        previous_paragraph = OxmlElement("w:p")
        paragraph._p.addprevious(previous_paragraph)
        previous = previous_paragraph
    ppr = previous.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        previous.insert(0, ppr)
    if ppr.find(qn("w:sectPr")) is not None:
        return
    sectpr = copy.deepcopy(template_sectpr)
    section_type = sectpr.find(qn("w:type"))
    if section_type is None:
        section_type = OxmlElement("w:type")
        sectpr.insert(0, section_type)
    section_type.set(qn("w:val"), "nextPage")
    for tag in (qn("w:headerReference"), qn("w:footerReference")):
        for element in list(sectpr.findall(tag)):
            sectpr.remove(element)
    ppr.append(sectpr)


def configure_sections(doc):
    if len(doc.sections) < 6:
        raise RuntimeError(f"Expected at least 6 sections after front-matter split, found {len(doc.sections)}")
    for index, section in enumerate(doc.sections):
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        section.first_page_header.is_linked_to_previous = False
        section.first_page_footer.is_linked_to_previous = False
        clear_container(section.header)
        clear_container(section.footer)
        clear_container(section.first_page_header)
        clear_container(section.first_page_footer)

        if index == 0:
            continue
        if index == 1:
            set_page_number_format(section, "lowerRoman", 1)
            footer = section.footer.add_paragraph()
            add_page_number(footer, center=True)
            continue

        set_page_number_format(section, "decimal", 1 if index == 2 else None)
        section.different_first_page_header_footer = True
        header = section.header.add_paragraph()
        add_page_number(header)
        first_footer = section.first_page_footer.add_paragraph()
        add_page_number(first_footer)


def remove_dynamic_toc(doc):
    paragraphs = doc.paragraphs
    toc_index = next(i for i, p in enumerate(paragraphs) if p.text.strip().upper() == "DAFTAR ISI")
    table_index = next(
        i for i, p in enumerate(paragraphs)
        if i > toc_index and p.text.strip().upper() == "DAFTAR TABEL"
    )
    start = paragraphs[toc_index]._p
    end = paragraphs[table_index]._p
    current = start.getnext()
    while current is not None and current is not end:
        nxt = current.getnext()
        current.getparent().remove(current)
        current = nxt

    marker = OxmlElement("w:p")
    ppr = OxmlElement("w:pPr")
    pstyle = OxmlElement("w:pStyle")
    pstyle.set(qn("w:val"), "StaticTOCMarker")
    ppr.append(pstyle)
    marker.append(ppr)
    start.addnext(marker)
    add_style(doc, "Static TOC Marker")


def prepare(input_path: Path, output_path: Path):
    doc = Document(input_path)
    ensure_styles(doc)
    add_style(doc, "Static TOC Marker")
    toc = next(p for p in doc.paragraphs if p.text.strip().upper() == "DAFTAR ISI")

    if len(doc.sections) == 5:
        add_section_break_before(toc, doc.sections[-1]._sectPr)
    doc.save(output_path)

    doc = Document(output_path)
    ensure_styles(doc)
    remove_dynamic_toc(doc)
    doc.save(output_path)

    doc = Document(output_path)
    configure_sections(doc)
    doc.save(output_path)


def normalize(text):
    text = text.replace("\u00a0", " ").replace("–", "-").replace("—", "-").replace("−", "-")
    return re.sub(r"\s+", " ", text).strip().casefold()


def to_roman(number):
    values = (
        (1000, "m"), (900, "cm"), (500, "d"), (400, "cd"),
        (100, "c"), (90, "xc"), (50, "l"), (40, "xl"),
        (10, "x"), (9, "ix"), (5, "v"), (4, "iv"), (1, "i"),
    )
    result = ""
    for value, symbol in values:
        while number >= value:
            result += symbol
            number -= value
    return result


def pdf_pages(pdf_path):
    pdf = fitz.open(pdf_path)
    return [normalize(page.get_text()) for page in pdf]


def find_page(pages, text, extra=None, start_page=1):
    query = normalize(text)
    extra_query = normalize(extra) if extra else None
    for page_number, page_text in enumerate(pages[start_page - 1:], start=start_page):
        if query in page_text and (extra_query is None or extra_query in page_text):
            return page_number
    short_query = query[:45]
    for page_number, page_text in enumerate(pages[start_page - 1:], start=start_page):
        if short_query and short_query in page_text and (
            extra_query is None or extra_query[:30] in page_text
        ):
            return page_number
    raise RuntimeError(f"Could not locate heading in rendered PDF: {text!r}")


def find_title_page(pages, text, start_page=1):
    query = normalize(text)
    candidates = []
    for page_number, page_text in enumerate(pages[start_page - 1:], start=start_page):
        if query in page_text:
            candidates.append((len(page_text), page_number))
    if not candidates:
        raise RuntimeError(f"Could not locate title in rendered PDF: {text!r}")
    return min(candidates)[1]


def collect_entries(doc, pages):
    paragraphs = doc.paragraphs
    bab1_page = find_page(pages, "BAB I", "Kualitas biji kopi merupakan")
    toc_page = find_title_page(pages, "DAFTAR ISI")
    entries = []

    for title in ("DAFTAR ISI", "DAFTAR TABEL", "DAFTAR GAMBAR"):
        physical_page = find_title_page(pages, title)
        entries.append((1, title, to_roman(physical_page - toc_page + 1)))

    index = 0
    while index < len(paragraphs):
        paragraph = paragraphs[index]
        text = paragraph.text.strip()
        chapter = re.fullmatch(r"BAB\s+([IVX]+)", text, re.I)
        if chapter and index + 1 < len(paragraphs):
            title = paragraphs[index + 1].text.strip()
            physical_page = find_page(pages, text, title, start_page=bab1_page)
            entries.append((1, f"{text.upper()} {title.upper()}", str(physical_page - bab1_page + 1)))
            index += 2
            continue

        if paragraph.style.name in ("Heading 2", "Heading 3", "Heading 4") and re.match(
            r"^\d+(?:\.\d+)+\s+", text
        ):
            physical_page = find_page(pages, text, start_page=bab1_page)
            level = {"Heading 2": 2, "Heading 3": 3, "Heading 4": 3}[paragraph.style.name]
            entries.append((level, text, str(physical_page - bab1_page + 1)))

        if text.upper() == "DAFTAR PUSTAKA":
            physical_page = find_page(pages, "DAFTAR PUSTAKA", start_page=bab1_page)
            entries.append((1, "DAFTAR PUSTAKA", str(physical_page - bab1_page + 1)))
        index += 1

    return entries


def insert_static_toc(doc, entries):
    ensure_styles(doc)
    paragraphs = doc.paragraphs
    toc_index = next(i for i, p in enumerate(paragraphs) if p.text.strip().upper() == "DAFTAR ISI")
    table_index = next(
        i for i, p in enumerate(paragraphs)
        if i > toc_index and p.text.strip().upper() == "DAFTAR TABEL"
    )
    title = paragraphs[toc_index]._p
    end = paragraphs[table_index]._p

    current = title.getnext()
    while current is not None and current is not end:
        nxt = current.getnext()
        current.getparent().remove(current)
        current = nxt

    anchor = title
    for level, text, page in entries:
        paragraph = OxmlElement("w:p")
        ppr = OxmlElement("w:pPr")
        pstyle = OxmlElement("w:pStyle")
        pstyle.set(qn("w:val"), f"StaticTOC{level}")
        ppr.append(pstyle)
        paragraph.append(ppr)

        text_run = OxmlElement("w:r")
        text_node = OxmlElement("w:t")
        text_node.text = text
        text_run.append(text_node)
        paragraph.append(text_run)

        tab_run = OxmlElement("w:r")
        tab_run.append(OxmlElement("w:tab"))
        paragraph.append(tab_run)

        page_run = OxmlElement("w:r")
        page_node = OxmlElement("w:t")
        page_node.text = page
        page_run.append(page_node)
        paragraph.append(page_run)

        anchor.addnext(paragraph)
        anchor = paragraph


def populate(input_path: Path, pdf_path: Path, output_path: Path):
    doc = Document(input_path)
    pages = pdf_pages(pdf_path)
    entries = collect_entries(doc, pages)
    insert_static_toc(doc, entries)
    doc.save(output_path)
    for level, text, page in entries:
        print(f"{level}: {text} -> {page}")


def main():
    parser = argparse.ArgumentParser(description="Build a deterministic static thesis TOC.")
    subparsers = parser.add_subparsers(dest="command", required=True)

    prepare_parser = subparsers.add_parser("prepare")
    prepare_parser.add_argument("input")
    prepare_parser.add_argument("output")

    populate_parser = subparsers.add_parser("populate")
    populate_parser.add_argument("input")
    populate_parser.add_argument("pdf")
    populate_parser.add_argument("output")

    args = parser.parse_args()
    if args.command == "prepare":
        prepare(Path(args.input), Path(args.output))
    else:
        populate(Path(args.input), Path(args.pdf), Path(args.output))


if __name__ == "__main__":
    main()
