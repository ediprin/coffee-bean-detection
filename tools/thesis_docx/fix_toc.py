#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

FONT = "Times New Roman"


def style_by_name(doc, name):
    return next(style for style in doc.styles if style.name == name)


def add_paragraph_style(doc, name, base="Normal"):
    if name in doc.styles:
        return style_by_name(doc, name)
    style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = style_by_name(doc, base)
    return style


def set_style_font(style, *, bold=False, size=12):
    style.font.name = FONT
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), FONT)


def configure_chapter_styles(doc):
    # Match the existing Heading 1 spacing exactly so TOC source changes do not
    # shift chapter pagination or alter the visible chapter layout.
    for name in ("Chapter Number USU", "Chapter Title USU"):
        style = add_paragraph_style(doc, name)
        set_style_font(style, bold=True)
        pf = style.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.line_spacing = 1.0
        pf.first_line_indent = Cm(0)
        pf.left_indent = Cm(0)
        pf.right_indent = Cm(0)
        pf.space_before = Pt(12)
        pf.space_after = Pt(6)
        pf.keep_with_next = False


def configure_toc_styles(doc):
    specs = {
        "TOC 1": (True, 0.0, 4),
        "TOC 2": (False, 0.75, 2),
        "TOC 3": (False, 1.50, 1),
    }
    for name, (bold, indent_cm, after_pt) in specs.items():
        style = add_paragraph_style(doc, name)
        set_style_font(style, bold=bold)
        pf = style.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.line_spacing = 1.0
        pf.first_line_indent = Cm(0)
        pf.left_indent = Cm(indent_cm)
        pf.right_indent = Cm(0)
        pf.space_before = Pt(0)
        pf.space_after = Pt(after_pt)
        try:
            pf.tab_stops.clear_all()
        except Exception:
            pass
        pf.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)


def add_tc_field(paragraph, text: str, level: int = 1, identifier: str = "C"):
    for fld in list(paragraph._p.findall(qn("w:fldSimple"))):
        instr = fld.get(qn("w:instr"), "")
        if instr.lstrip().startswith("TC "):
            paragraph._p.remove(fld)

    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), f'TC "{text}" \\f {identifier} \\l {level}')
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    vanish = OxmlElement("w:vanish")
    rpr.append(vanish)
    run.append(rpr)
    fld.append(run)
    paragraph._p.append(fld)


def patch_main_toc_field(doc):
    found = False
    for instr in doc.element.iter(qn("w:instrText")):
        text = instr.text or ""
        if re.search(r"\bTOC\b", text) and ('\\o "1-3"' in text or "\\o '1-3'" in text):
            # Level 1 comes from explicit TC entries so BAB number + title form one line.
            # Levels 2-3 still come directly from the numbered Markdown headings.
            instr.text = ' TOC \\f C \\o "2-3" \\h \\z \\u '
            found = True
            break
    if not found:
        raise RuntimeError("Main TOC field was not found")


def force_run_font(paragraph, bold=None):
    for run in paragraph.runs:
        run.font.name = FONT
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
        rfonts = run._r.get_or_add_rPr().get_or_add_rFonts()
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            rfonts.set(qn(f"w:{attr}"), FONT)
        if bold is not None:
            run.bold = bold


def patch_chapters(doc):
    paragraphs = doc.paragraphs
    chapter_entries = []
    i = 0
    while i < len(paragraphs) - 1:
        p = paragraphs[i]
        match = re.fullmatch(r"BAB\s+([IVX]+)", p.text.strip(), re.I)
        if not match:
            i += 1
            continue

        title_p = paragraphs[i + 1]
        title = title_p.text.strip()
        if not title or re.match(r"^\d+(?:\.\d+)+\s", title):
            raise RuntimeError(f"Chapter title missing after {p.text!r}")

        roman = match.group(1).upper()
        toc_text = f"BAB {roman} {title.upper()}"
        p.style = style_by_name(doc, "Chapter Number USU")
        title_p.style = style_by_name(doc, "Chapter Title USU")
        p.paragraph_format.page_break_before = True
        title_p.paragraph_format.page_break_before = False
        force_run_font(p, bold=True)
        force_run_font(title_p, bold=True)
        add_tc_field(p, toc_text, 1)
        chapter_entries.append(toc_text)
        i += 2

    if len(chapter_entries) < 3:
        raise RuntimeError(f"Expected at least 3 chapter entries, found {chapter_entries}")

    for p in paragraphs:
        if p.text.strip().upper() == "DAFTAR PUSTAKA":
            p.style = style_by_name(doc, "Chapter Title USU")
            p.paragraph_format.page_break_before = True
            force_run_font(p, bold=True)
            add_tc_field(p, "DAFTAR PUSTAKA", 1)
            break
    else:
        raise RuntimeError("DAFTAR PUSTAKA heading not found")

    return chapter_entries


def set_update_fields(doc):
    settings = doc.settings._element
    element = settings.find(qn("w:updateFields"))
    if element is None:
        element = OxmlElement("w:updateFields")
        settings.append(element)
    element.set(qn("w:val"), "true")


def fix_toc(input_path: Path, output_path: Path):
    doc = Document(input_path)
    configure_chapter_styles(doc)
    configure_toc_styles(doc)
    entries = patch_chapters(doc)
    patch_main_toc_field(doc)
    set_update_fields(doc)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print("TOC chapter entries:")
    for entry in entries:
        print(f"- {entry}")
    print("- DAFTAR PUSTAKA")


def main():
    parser = argparse.ArgumentParser(description="Format USU proposal TOC chapter entries.")
    parser.add_argument("input")
    parser.add_argument("--output", default=None)
    args = parser.parse_args()
    source = Path(args.input)
    target = Path(args.output) if args.output else source
    fix_toc(source, target)


if __name__ == "__main__":
    main()
