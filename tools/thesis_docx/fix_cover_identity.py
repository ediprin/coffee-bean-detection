#!/usr/bin/env python3
from pathlib import Path
import argparse

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

FONT = "Times New Roman"
COVER_LINES = [
    "PROGRAM STUDI S2 TEKNIK INFORMATIKA",
    "FAKULTAS ILMU KOMPUTER DAN TEKNOLOGI INFORMASI",
    "UNIVERSITAS SUMATERA UTARA",
    "MEDAN",
    "2026",
]


def format_cover_line(paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = FONT
    run.font.size = Pt(12)


def apply_cover_identity(input_path: Path, output_path: Path) -> None:
    doc = Document(input_path)
    paragraphs = list(doc.paragraphs)

    # Fresh build_proposal.py output still contains the legacy SPs block.
    legacy_index = next(
        (i for i, p in enumerate(paragraphs) if p.text.strip().upper() == "SEKOLAH PASCASARJANA"),
        None,
    )

    if legacy_index is not None:
        if legacy_index < 1 or legacy_index + 2 >= len(paragraphs):
            raise RuntimeError("Unexpected cover structure around SEKOLAH PASCASARJANA")

        target = [
            paragraphs[legacy_index - 1],
            paragraphs[legacy_index],
            paragraphs[legacy_index + 1],
            paragraphs[legacy_index + 2],
        ]
        for paragraph, text in zip(target, COVER_LINES[:4]):
            format_cover_line(paragraph, text)

        # Insert the final year line immediately after MEDAN.
        year_paragraph = doc.add_paragraph()
        target[-1]._p.addnext(year_paragraph._p)
        format_cover_line(year_paragraph, COVER_LINES[4])
    else:
        # Idempotent validation for an already-correct file.
        texts = [p.text.strip().upper() for p in paragraphs]
        try:
            start = texts.index(COVER_LINES[0])
        except ValueError as exc:
            raise RuntimeError("Could not locate cover institutional block") from exc
        actual = texts[start : start + len(COVER_LINES)]
        if actual != COVER_LINES:
            raise RuntimeError(f"Unexpected existing cover block: {actual}")
        for paragraph, text in zip(paragraphs[start : start + len(COVER_LINES)], COVER_LINES):
            format_cover_line(paragraph, text)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print("Cover identity set to Fasilkom-TI USU, Medan, 2026.")


def main() -> None:
    parser = argparse.ArgumentParser(description="Set the exact institutional identity block on the thesis cover.")
    parser.add_argument("input")
    parser.add_argument("--output", default=None)
    args = parser.parse_args()
    source = Path(args.input)
    target = Path(args.output) if args.output else source
    apply_cover_identity(source, target)


if __name__ == "__main__":
    main()
