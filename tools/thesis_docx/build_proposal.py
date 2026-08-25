#!/usr/bin/env python3
from __future__ import annotations

import argparse
import base64
import copy
import os
import re
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

LOGO_B64_PATH = Path(__file__).with_name("usu_logo.jpg.b64")
ROMAN = {"I": 1, "II": 2, "III": 3, "IV": 4, "V": 5, "VI": 6, "VII": 7, "VIII": 8, "IX": 9, "X": 10}
FORMAL_PROPOSAL_DIR = Path("docs/thesis/proposal")


def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_field(paragraph, instruction: str, result_text: str = ""):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text_run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = result_text
    text_run.append(text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text_run, end])


def add_page_number(paragraph, fmt="decimal"):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if fmt == "lowerRoman" else WD_ALIGN_PARAGRAPH.RIGHT
    add_field(paragraph, " PAGE ", "1")


def set_page_number_format(section, fmt="decimal", start=None):
    sect_pr = section._sectPr
    pg = sect_pr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType")
        sect_pr.append(pg)
    pg.set(qn("w:fmt"), fmt)
    if start is not None:
        pg.set(qn("w:start"), str(start))
    elif qn("w:start") in pg.attrib:
        del pg.attrib[qn("w:start")]


def set_update_fields(doc):
    settings = doc.settings._element
    element = settings.find(qn("w:updateFields"))
    if element is None:
        element = OxmlElement("w:updateFields")
        settings.append(element)
    element.set(qn("w:val"), "true")


def style_by_name(doc, name):
    return next(style for style in doc.styles if style.name == name)


def add_style(doc, name, base="Normal"):
    if name in doc.styles:
        return style_by_name(doc, name)
    style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = style_by_name(doc, base)
    return style


def set_font(style, size=12, bold=None, italic=None):
    style.font.name = "Times New Roman"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor(0, 0, 0)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic


def configure_styles(doc):
    for name in ("Normal", "Body Text", "First Paragraph"):
        if name in doc.styles:
            style = style_by_name(doc, name)
            set_font(style, 12)
            pf = style.paragraph_format
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf.line_spacing = 1.5
            pf.first_line_indent = Cm(1.27)
            pf.space_after = Pt(0)
            pf.space_before = Pt(0)

    for level in (1, 2, 3, 4):
        style = style_by_name(doc, f"Heading {level}")
        set_font(style, 12, bold=True)
        pf = style.paragraph_format
        pf.line_spacing = 1.0
        pf.first_line_indent = Cm(0)
        pf.space_after = Pt(6)
        pf.space_before = Pt(12)
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT

    caption_table = add_style(doc, "Caption Table")
    set_font(caption_table, 12, bold=True)
    caption_table.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_table.paragraph_format.line_spacing = 1.0
    caption_table.paragraph_format.first_line_indent = Cm(0)

    caption_figure = add_style(doc, "Caption Figure")
    set_font(caption_figure, 12, bold=False)
    caption_figure.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_figure.paragraph_format.line_spacing = 1.0
    caption_figure.paragraph_format.first_line_indent = Cm(0)

    front = add_style(doc, "Front Matter Title")
    set_font(front, 12, bold=True)
    front.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    front.paragraph_format.line_spacing = 1.0
    front.paragraph_format.first_line_indent = Cm(0)

    bibliography = add_style(doc, "Bibliography USU")
    set_font(bibliography, 12)
    bibliography.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    bibliography.paragraph_format.line_spacing = 1.0
    bibliography.paragraph_format.left_indent = Cm(1.27)
    bibliography.paragraph_format.first_line_indent = Cm(-1.27)
    bibliography.paragraph_format.space_after = Pt(12)


def configure_section_layout(section):
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)


def normalize_headings_and_captions(doc):
    in_bibliography = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        match = re.match(r"^BAB\s+([IVX]+)$", text, re.I)
        if match:
            paragraph.style = style_by_name(doc, "Heading 1")
            paragraph.paragraph_format.page_break_before = True
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)
            continue

        if text.upper() == "DAFTAR PUSTAKA":
            in_bibliography = True
            paragraph.style = style_by_name(doc, "Heading 1")
            paragraph.paragraph_format.page_break_before = True
            continue

        if in_bibliography and text and paragraph.style.name != "Heading 1":
            paragraph.style = style_by_name(doc, "Bibliography USU")

        table_match = re.match(r"^(Tabel)\s+(\d+\.\d+)\s*:?[.]?\s*(.*)$", text, re.I)
        if table_match:
            paragraph.text = f"Tabel {table_match.group(2)}: {table_match.group(3)}".rstrip()
            paragraph.style = style_by_name(doc, "Caption Table")
            continue

        figure_match = re.match(r"^(Gambar)\s+(\d+\.\d+)\s*:?[.]?\s*(.*)$", text, re.I)
        if figure_match:
            paragraph.text = f"Gambar {figure_match.group(2)}. {figure_match.group(3)}".rstrip()
            paragraph.style = style_by_name(doc, "Caption Figure")
            continue

        if paragraph.style.name in ("Heading 2", "Heading 3", "Heading 4"):
            paragraph.paragraph_format.first_line_indent = Cm(0)
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)

    # Markdown proposal uses two consecutive H1s, e.g. "BAB III" then "METODE PENELITIAN".
    # Only the first H1 should force the chapter page break.
    for index, paragraph in enumerate(doc.paragraphs[:-1]):
        if re.match(r"^BAB\s+[IVX]+$", paragraph.text.strip(), re.I):
            next_paragraph = doc.paragraphs[index + 1]
            if next_paragraph.style.name == "Heading 1":
                next_paragraph.paragraph_format.page_break_before = False


def format_tables(doc):
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        if table.rows:
            set_repeat_table_header(table.rows[0])
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell, top=60, start=80, bottom=60, end=80)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.paragraph_format.first_line_indent = Cm(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)
        if table.rows:
            for cell in table.rows[0].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def number_equations(doc):
    chapter = 0
    counters = {}
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:p"):
            continue
        text = "".join(child.xpath(".//w:t/text()")).strip()
        chapter_match = re.match(r"^BAB\s+([IVX]+)$", text, re.I)
        if chapter_match:
            chapter = ROMAN.get(chapter_match.group(1).upper(), chapter)
        math_nodes = child.findall(qn("m:oMathPara"))
        if not math_nodes:
            continue
        counters[chapter] = counters.get(chapter, 0) + 1
        number = f"({chapter}.{counters[chapter]})" if chapter else f"({counters[chapter]})"
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        remove_table_borders(table)
        widths = (Cm(1.0), Cm(11.5), Cm(1.5))
        for cell, width in zip(table.rows[0].cells, widths):
            cell.width = width
            set_cell_margins(cell, 0, 0, 0, 0)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        center = table.rows[0].cells[1].paragraphs[0]
        center.alignment = WD_ALIGN_PARAGRAPH.CENTER
        center.paragraph_format.line_spacing = 1.0
        center.paragraph_format.first_line_indent = Cm(0)
        for node in list(center._p):
            if node.tag != qn("w:pPr"):
                center._p.remove(node)
        for math_node in math_nodes:
            center._p.append(copy.deepcopy(math_node))

        right = table.rows[0].cells[2].paragraphs[0]
        right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right.paragraph_format.line_spacing = 1.0
        right.paragraph_format.first_line_indent = Cm(0)
        run = right.add_run(number)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

        child.addnext(table._tbl)
        body.remove(child)


def make_break_before(paragraph, template_sectpr):
    previous = paragraph._p.getprevious()
    if previous is None:
        return
    if previous.tag != qn("w:p"):
        previous_paragraph = OxmlElement("w:p")
        paragraph._p.addprevious(previous_paragraph)
        previous = previous_paragraph
    p_pr = previous.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        previous.insert(0, p_pr)
    old = p_pr.find(qn("w:sectPr"))
    if old is not None:
        p_pr.remove(old)
    section_properties = copy.deepcopy(template_sectpr)
    section_type = section_properties.find(qn("w:type"))
    if section_type is None:
        section_type = OxmlElement("w:type")
        section_properties.insert(0, section_type)
    section_type.set(qn("w:val"), "nextPage")
    for tag in (qn("w:headerReference"), qn("w:footerReference")):
        for element in list(section_properties.findall(tag)):
            section_properties.remove(element)
    p_pr.append(section_properties)


def clear_container(container):
    for paragraph in container.paragraphs:
        paragraph._element.getparent().remove(paragraph._element)
    for table in container.tables:
        table._element.getparent().remove(table._element)


def configure_headers_footers(doc):
    for section in doc.sections:
        configure_section_layout(section)

    # Expected sections: cover, front matter, BAB I, BAB II, BAB III, bibliography.
    for index, section in enumerate(doc.sections):
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        section.first_page_header.is_linked_to_previous = False
        section.first_page_footer.is_linked_to_previous = False
        clear_container(section.header)
        clear_container(section.footer)
        clear_container(section.first_page_header)
        clear_container(section.first_page_footer)

        if index == 0:
            continue
        if index == 1:
            set_page_number_format(section, "lowerRoman", 1)
            paragraph = section.footer.add_paragraph()
            add_page_number(paragraph, "lowerRoman")
            continue

        if index == 2:
            set_page_number_format(section, "decimal", 1)
        else:
            set_page_number_format(section, "decimal", None)
        section.different_first_page_header_footer = True

        header = section.header.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_field(header, " PAGE ", "1")
        first_footer = section.first_page_footer.add_paragraph()
        first_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_field(first_footer, " PAGE ", "1")
        for paragraph in (header, first_footer):
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)


def insert_front_matter(doc, title, student, nim, prodi, year, label):
    logo_path = Path(tempfile.gettempdir()) / "usu_logo.jpg"
    logo_path.write_bytes(base64.b64decode(LOGO_B64_PATH.read_text(encoding="ascii").strip()))
    created = []

    def paragraph(text="", bold=False, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=0):
        item = doc.add_paragraph()
        item.alignment = align
        item.paragraph_format.first_line_indent = Cm(0)
        item.paragraph_format.line_spacing = 1.0
        item.paragraph_format.space_before = Pt(space_before)
        item.paragraph_format.space_after = Pt(space_after)
        run = item.add_run(text)
        run.bold = bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        created.append(item)
        return item

    paragraph(title.upper(), True, 14, space_after=18)
    paragraph(label.upper(), True, 14, space_after=30)
    paragraph("Oleh", True, 12, space_after=12)
    paragraph(student.upper(), True, 12)
    paragraph(nim.upper(), True, 12, space_after=18)

    image_paragraph = doc.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.first_line_indent = Cm(0)
    image_paragraph.add_run().add_picture(str(logo_path), width=Cm(4.31), height=Cm(4.31))
    created.append(image_paragraph)

    paragraph(prodi.upper(), True, 12, space_before=24)
    paragraph("SEKOLAH PASCASARJANA", True, 12)
    paragraph("UNIVERSITAS SUMATERA UTARA", True, 12)
    paragraph(str(year), True, 12)

    marker = doc.add_paragraph("__FRONT_MATTER__")
    marker.paragraph_format.page_break_before = True
    created.append(marker)

    toc_title = doc.add_paragraph()
    toc_title.style = style_by_name(doc, "Front Matter Title")
    toc_title.add_run("DAFTAR ISI")
    created.append(toc_title)
    toc = doc.add_paragraph()
    toc.paragraph_format.first_line_indent = Cm(0)
    add_field(toc, ' TOC \\o "1-3" \\h \\z \\u ', "Daftar isi diperbarui saat dokumen dibuka.")
    created.append(toc)

    tables_title = doc.add_paragraph()
    tables_title.paragraph_format.page_break_before = True
    tables_title.style = style_by_name(doc, "Front Matter Title")
    tables_title.add_run("DAFTAR TABEL")
    created.append(tables_title)
    tables_toc = doc.add_paragraph()
    tables_toc.paragraph_format.first_line_indent = Cm(0)
    add_field(tables_toc, ' TOC \\h \\z \\t "Caption Table,1" ', "Daftar tabel diperbarui saat dokumen dibuka.")
    created.append(tables_toc)

    figures_title = doc.add_paragraph()
    figures_title.paragraph_format.page_break_before = True
    figures_title.style = style_by_name(doc, "Front Matter Title")
    figures_title.add_run("DAFTAR GAMBAR")
    created.append(figures_title)
    figures_toc = doc.add_paragraph()
    figures_toc.paragraph_format.first_line_indent = Cm(0)
    add_field(figures_toc, ' TOC \\h \\z \\t "Caption Figure,1" ', "Daftar gambar diperbarui saat dokumen dibuka.")
    created.append(figures_toc)

    body = doc._element.body
    for item in reversed(created):
        body.remove(item._p)
        body.insert(0, item._p)

    for item in doc.paragraphs:
        if item.text.strip() == "__FRONT_MATTER__":
            item.text = ""
            item.paragraph_format.page_break_before = True


def extract_title(repo: Path):
    readme = repo / "README.md"
    if readme.exists():
        text = readme.read_text(encoding="utf-8")
        match = re.search(r"##\s+Judul kerja.*?\*\*(.+?)\*\*", text, re.S | re.I)
        if match:
            return " ".join(match.group(1).split())
    return "Analisis dan Optimasi Prapemrosesan Citra Berbasis Frekuensi-Angular pada YOLO26 untuk Deteksi Fine-Grained Cacat Biji Kopi"


def build(repo: Path, output: Path, student, nim, prodi, year, label):
    proposal_dir = repo / FORMAL_PROPOSAL_DIR
    sources = [
        proposal_dir / "BAB_I_PENDAHULUAN.md",
        proposal_dir / "BAB_II_TINJAUAN_PUSTAKA.md",
        proposal_dir / "BAB_III_METODOLOGI_PENELITIAN.md",
        proposal_dir / "DAFTAR_PUSTAKA.md",
    ]
    missing = [str(path) for path in sources if not path.exists()]
    if missing:
        raise SystemExit("Missing formal proposal source files: " + ", ".join(missing))

    combined = "\n\n".join(path.read_text(encoding="utf-8") for path in sources)
    title = extract_title(repo)
    output.parent.mkdir(parents=True, exist_ok=True)

    with tempfile.TemporaryDirectory() as directory:
        directory = Path(directory)
        markdown = directory / "proposal.md"
        raw_docx = directory / "raw.docx"
        sectioned_docx = directory / "sectioned.docx"
        markdown.write_text(combined, encoding="utf-8")

        subprocess.run(
            [
                "pandoc",
                "--from=markdown+tex_math_dollars+tex_math_single_backslash+pipe_tables",
                "--to=docx",
                f"--resource-path={proposal_dir}{os.pathsep}{repo}",
                str(markdown),
                "-o",
                str(raw_docx),
            ],
            check=True,
        )

        doc = Document(raw_docx)
        configure_styles(doc)
        for section in doc.sections:
            configure_section_layout(section)
        normalize_headings_and_captions(doc)
        format_tables(doc)
        number_equations(doc)
        insert_front_matter(doc, title, student, nim, prodi, year, label)
        set_update_fields(doc)
        doc.save(sectioned_docx)

        doc = Document(sectioned_docx)
        template_sectpr = copy.deepcopy(doc.sections[-1]._sectPr)
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if re.match(r"^BAB\s+[IVX]+$", text, re.I) or text.upper() == "DAFTAR PUSTAKA":
                make_break_before(paragraph, template_sectpr)
        doc.save(output)

        doc = Document(output)
        configure_headers_footers(doc)
        set_update_fields(doc)
        doc.save(output)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--repo", default=".")
    parser.add_argument("--output", required=True)
    parser.add_argument("--student", default=os.getenv("THESIS_STUDENT", "NAMA MAHASISWA"))
    parser.add_argument("--nim", default=os.getenv("THESIS_NIM", "NIM / SINGKATAN PRODI"))
    parser.add_argument("--prodi", default=os.getenv("THESIS_PRODI", "NAMA PROGRAM STUDI"))
    parser.add_argument("--year", default=os.getenv("THESIS_YEAR", "2026"))
    parser.add_argument("--label", default=os.getenv("THESIS_LABEL", "TESIS"))
    args = parser.parse_args()
    build(
        Path(args.repo).resolve(),
        Path(args.output).resolve(),
        args.student,
        args.nim,
        args.prodi,
        args.year,
        args.label,
    )


if __name__ == "__main__":
    main()
